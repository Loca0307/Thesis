LINK NUMBER 1
Not enough lines

LINK NUMBER 2
Not enough lines

LINK NUMBER 3

File path: Libman.Audit/LibmanPackage.cs
"using System;
using System.Collections.Generic;
using System.IO;
using System.Linq;
using System.Net.Http;
using System.Text;
using System.Text.Json;
using System.Text.Json.Serialization;
using System.Threading.Tasks;
using Microsoft.Build.Framework;
using Microsoft.Build.Utilities;

using Task = Microsoft.Build.Utilities.Task;

namespace Libman.Audit
{
    public class LibmanAuditTask : Task
    {
        [Required]
        public string LibmanJsonPath { get; set; }

        [Output]
        public ITaskItem[] VulnerablePackages { get; private set; }

        private readonly HttpClient _httpClient;
        private const string SonatypeApiBaseUrl = ""https://ossindex.sonatype.org/api/v3/component-report"";
        private static readonly JsonSerializerOptions _jsonOptions = new JsonSerializerOptions
        {
            PropertyNameCaseInsensitive = true,
            PropertyNamingPolicy = JsonNamingPolicy.CamelCase
        };

        public LibmanAuditTask()
        {
            _httpClient = new HttpClient();
            // initialize non-null values
            LibmanJsonPath = """";
            VulnerablePackages = [];
        }

        public override bool Execute()
        {
            try
            {
                Log.LogMessage(MessageImportance.Normal, ""Starting Libman audit task..."");

                if (!File.Exists(LibmanJsonPath))
                {
                    Log.LogError($""Libman.json file not found at: {LibmanJsonPath}"");
                    return false;
                }

                string jsonContent = File.ReadAllText(LibmanJsonPath);
                List<LibmanPackage> libmanPackages = ParseLibmanJson(jsonContent);

                if (libmanPackages.Count == 0)
                {
                    Log.LogMessage(MessageImportance.Normal, ""No packages found in libman.json"");
                    VulnerablePackages = new TaskItem[0];
                    return true;
                }

                List<VulnerablePackage> vulnerablePackages = AuditPackagesAsync(libmanPackages).GetAwaiter().GetResult();
                VulnerablePackages = ConvertToTaskItems(vulnerablePackages);

                if (vulnerablePackages.Count > 0)
                {
                    Log.LogWarning($""Found {vulnerablePackages.Count} vulnerable packages in libman.json"");
                    foreach (VulnerablePackage package in vulnerablePackages)
                    {
                        Log.LogWarning($""Vulnerable package: {package.Name} {package.Version}, Vulnerability count: {package.VulnerabilityCount}"");
                    }
                }
                else
                {
                    Log.LogMessage(MessageImportance.Normal, ""No vulnerable packages found"");
                }

                return true;
            }
            catch (Exception ex)
            {
                Log.LogErrorFromException(ex);
                return false;
            }
        }

        private List<LibmanPackage> ParseLibmanJson(string jsonContent)
        {
            List<LibmanPackage> packages = new List<LibmanPackage>();

            try
            {
                using (JsonDocument doc = JsonDocument.Parse(jsonContent))
                {
                    JsonElement root = doc.RootElement;

                    if (!root.TryGetProperty(""libraries"", out JsonElement librariesElement) ||
                        librariesElement.ValueKind != JsonValueKind.Array)
                    {
                        Log.LogMessage(MessageImportance.Normal, ""No libraries found in libman.json"");
                        return packages;
                    }

                    foreach (JsonElement library in librariesElement.EnumerateArray())
                    {
                        if (!library.TryGetProperty(""provider"", out JsonElement providerElement) ||
                            !library.TryGetProperty(""library"", out JsonElement nameElement))
                        {
                            continue;
                        }

                        string provider = providerElement.GetString() ?? string.Empty;
                        string name = nameElement.GetString() ?? string.Empty;

                        if (string.IsNullOrEmpty(provider) || string.IsNullOrEmpty(name))
                        {
                            continue;
                        }

                        // Parse package name and version (format varies by provider)
                        string packageName;
                        string packageVersion;

                        if (name.Contains(""@""))
                        {
                            string[] parts = name.Split(new[] { '@' }, 2);
                            packageName = parts[0];
                            packageVersion = parts[1];
                        }
                        else
                        {
                            // If no version is specified, use the name as-is and leave version empty
                            packageName = name;
                            packageVersion = string.Empty;
                        }

                        packages.Add(new LibmanPackage
                        {
                            Name = packageName,
                            Version = packageVersion,
                            Provider = provider
                        });

                        Log.LogMessage(MessageImportance.Low, $""Found package: {packageName} {packageVersion} (Provider: {provider})"");
                    }
                }
            }
            catch (JsonException ex)
            {
                Log.LogError($""Failed to parse libman.json: {ex.Message}"");
            }

            return packages;
        }

        private async Task<List<VulnerablePackage>> AuditPackagesAsync(List<LibmanPackage> packages)
        {
            List<VulnerablePackage> vulnerablePackages = new List<VulnerablePackage>();

            try
            {
                // Group packages in batches to avoid large requests
                for (int i = 0; i < packages.Count; i += 20)
                {
                    List<LibmanPackage> batch = packages.Skip(i).Take(20).ToList();
                    List<string> components = new List<string>();

                    foreach (LibmanPackage package in batch)
                    {
                        string packageId = GetPackageCoordinates(package);
                        if (!string.IsNullOrEmpty(packageId))
                        {
                            components.Add(packageId);
                        }
                    }

                    if (components.Count > 0)
                    {
                        SonatypeRequest requestData = new SonatypeRequest
                        {
                            Coordinates = components.ToArray()
                        };

                        StringContent content = new StringContent(
                            JsonSerializer.Serialize(requestData, _jsonOptions),
                            Encoding.UTF8,
                            ""application/json"");

                        HttpResponseMessage response = await _httpClient.PostAsync(SonatypeApiBaseUrl, content);

                        if (response.IsSuccessStatusCode)
                        {
                            string responseContent = await response.Content.ReadAsStringAsync();
                            List<SonatypeResult>? results = JsonSerializer.Deserialize<List<SonatypeResult>>(responseContent, _jsonOptions);

                            if (results != null)
                            {
                                foreach (SonatypeResult result in results)
                                {
                                    if (result.Vulnerabilities != null && result.Vulnerabilities.Count > 0)
                                    {
                                        LibmanPackage? package = batch.FirstOrDefault(p => GetPackageCoordinates(p) == result.Coordinates);
                                        if (package != null)
                                        {
                                            vulnerablePackages.Add(new VulnerablePackage
                                            {
                                                Name = package.Name,
                                                Version = package.Version,
                                                Provider = package.Provider,
                                                VulnerabilityCount = result.Vulnerabilities.Count,
                                                Description = string.Join(""; "", result.Vulnerabilities.Select(v => v.Title))
                                            });
                                        }
                                    }
                                }
                            }
                        }
                        else
                        {
                            Log.LogWarning($""Failed to get vulnerability data: {response.StatusCode} {await response.Content.ReadAsStringAsync()}"");
                        }
                    }
                }
            }
            catch (Exception ex)
            {
                Log.LogWarning($""Error checking for vulnerabilities: {ex.Message}"");
            }

            return vulnerablePackages;
        }

        private string GetPackageCoordinates(LibmanPackage package)
        {
            // Map libman providers to Sonatype coordinate formats
            switch (package.Provider.ToLowerInvariant())
            {
                case ""cdnjs"":
                    return $""pkg:npm/{package.Name}@{package.Version}"";
                case ""unpkg"":
                    return $""pkg:npm/{package.Name}@{package.Version}"";
                case ""jsdelivr"":
                    return $""pkg:npm/{package.Name}@{package.Version}"";
                default:
                    Log.LogWarning($""Unsupported provider: {package.Provider}"");
                    return """";
            }
        }

        private ITaskItem[] ConvertToTaskItems(List<VulnerablePackage> vulnerablePackages)
        {
            List<TaskItem> taskItems = new List<TaskItem>();

            foreach (VulnerablePackage package in vulnerablePackages)
            {
                TaskItem taskItem = new TaskItem(package.Name);
                taskItem.SetMetadata(""Version"", package.Version);
                taskItem.SetMetadata(""Provider"", package.Provider);
                taskItem.SetMetadata(""VulnerabilityCount"", package.VulnerabilityCount.ToString());
                taskItem.SetMetadata(""Description"", package.Description);
                taskItems.Add(taskItem);
            }

            return taskItems.ToArray();
        }
    }
}"

LINK NUMBER 4
Not enough lines

LINK NUMBER 5
Not enough lines

LINK NUMBER 6
Not enough lines

LINK NUMBER 7
Not enough lines

LINK NUMBER 8

File path: src/test/java/com/acme/basic/FibonacciControllerIT.java
"package com.acme.basic;

import org.junit.jupiter.api.Test;
import org.springframework.boot.test.context.SpringBootTest;
import org.springframework.boot.test.web.server.LocalServerPort;
import org.springframework.http.ResponseEntity;
import org.springframework.web.client.RestTemplate;

import static org.assertj.core.api.Assertions.assertThat;

@SpringBootTest(webEnvironment = SpringBootTest.WebEnvironment.RANDOM_PORT)
public class FibonacciControllerIT {

    @LocalServerPort
    private int port;

    private final RestTemplate restTemplate = new RestTemplate();

    @Test
    public void testFibonacciEndpoint() {
        String url = ""http://localhost:"" + port + ""/fibonacci?n=5"";
        ResponseEntity<String> response = restTemplate.getForEntity(url, String.class);

        assertThat(response.getStatusCode().is2xxSuccessful()).isTrue();
        assertThat(response.getBody()).isEqualTo(""5"");
    }
}"

LINK NUMBER 9
Not enough lines

LINK NUMBER 10
Not enough lines

LINK NUMBER 11

File path: 03.03.2025.claude-auto-app-copilot-mod.py
"import requests
import pandas as pd
import re
import os
from bs4 import BeautifulSoup
from datetime import datetime, timedelta
import pickle
from google.colab import drive
from google_auth_oauthlib.flow import InstalledAppFlow
from googleapiclient.discovery import build
from docx import Document
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer

# Initialization
def initialize():
    drive.mount('/content/drive')
    nltk.download('punkt')
    nltk.download('stopwords')

# Job Retriever Class
class JobRetriever:
    def __init__(self):
        self.jobs_df = pd.DataFrame(columns=['Title', 'Company', 'Location', 'Description', 'URL', 'Date_Posted', 'Keywords_Match'])

    def search_indeed(self, keywords, location, pages=5):
        base_url = ""https://www.indeed.com/jobs""
        all_jobs = []

        for page in range(pages):
            params = {'q': keywords, 'l': location, 'start': page * 10}
            headers = {'User-Agent': 'Mozilla/5.0'}

            try:
                response = requests.get(base_url, params=params, headers=headers)
                soup = BeautifulSoup(response.text, 'html.parser')
                job_cards = soup.find_all('div', class_='jobsearch-SerpJobCard')

                for card in job_cards:
                    job_title_elem = card.find('a', class_='jobtitle')
                    company_elem = card.find('span', class_='company')
                    location_elem = card.find('div', class_='recJobLoc')
                    description_elem = card.find('div', class_='summary')

                    if job_title_elem and company_elem:
                        job_title = job_title_elem.text.strip()
                        company = company_elem.text.strip()
                        location = location_elem['data-rc-loc'] if location_elem else ""N/A""
                        description = description_elem.text.strip() if description_elem else ""N/A""
                        url = ""https://www.indeed.com"" + job_title_elem['href']

                        keywords_list = keywords.lower().split()
                        match_score = sum(1 for keyword in keywords_list if keyword.lower() in (job_title.lower() + "" "" + description.lower()))

                        job_data = {'Title': job_title, 'Company': company, 'Location': location, 'Description': description, 'URL': url, 'Date_Posted': datetime.now().strftime(""%Y-%m-%d""), 'Keywords_Match': match_score}
                        all_jobs.append(job_data)

            except requests.RequestException as e:
                print(f""Error scraping Indeed page {page}: {str(e)}"")

        self.jobs_df = pd.concat([self.jobs_df, pd.DataFrame(all_jobs)], ignore_index=True)

    def filter_jobs(self, min_keywords_match=2):
        return self.jobs_df[self.jobs_df['Keywords_Match'] >= min_keywords_match].sort_values('Keywords_Match', ascending=False)

    def save_jobs(self, filename='job_listings.csv'):
        path = '/content/drive/My Drive/' + filename
        self.jobs_df.to_csv(path, index=False)
        print(f""Saved {len(self.jobs_df)} jobs to {path}"")
        return path

# Document Customizer Class
class DocumentCustomizer:
    def __init__(self, resume_path, cover_letter_path):
        self.resume_path = resume_path
        self.cover_letter_path = cover_letter_path
        self.resume_doc = self.load_document(resume_path)
        self.cover_letter_doc = self.load_document(cover_letter_path)

    def load_document(self, path):
        try:
            return Document(path)
        except Exception as e:
            print(f""Error loading document from {path}: {str(e)}"")
            return None

    def extract_job_keywords(self, job_description):
        tokens = word_tokenize(job_description.lower())
        filtered_tokens = [word for word in tokens if word.isalnum() and word not in stopwords.words('english')]
        vectorizer = TfidfVectorizer(max_features=20)
        tfidf_matrix = vectorizer.fit_transform([' '.join(filtered_tokens)])
        feature_names = vectorizer.get_feature_names_out()
        word_scores = [(word, tfidf_matrix[0, i]) for i, word in enumerate(feature_names)]
        word_scores.sort(key=lambda x: x[1], reverse=True)
        return [word for word, score in word_scores[:10]]

    def customize_resume(self, job_title, company_name, job_description):
        if not self.resume_doc:
            return None

        custom_resume = Document()
        keywords = self.extract_job_keywords(job_description)
        print(f""Keywords extracted: {keywords}"")

        for para in self.resume_doc.paragraphs:
            if ""[OBJECTIVE]"" in para.text:
                custom_text = para.text.replace(""[OBJECTIVE]"", f""Experienced professional seeking the {job_title} position at {company_name}, bringing expertise in {', '.join(keywords[:3])}."")
                custom_resume.add_paragraph(custom_text, para.style)
            else:
                text = para.text
                for keyword in keywords:
                    if keyword.lower() in text.lower() and len(keyword) > 3:
                        text = re.sub(re.escape(keyword), f""**{keyword}**"", text, flags=re.IGNORECASE)
                custom_resume.add_paragraph(text, para.style)

        filename = f""Custom_Resume_{company_name}_{datetime.now().strftime('%Y%m%d')}.docx""
        save_path = f""/content/drive/My Drive/{filename}""
        custom_resume.save(save_path)
        print(f""Customized resume saved to {save_path}"")
        return save_path

# Initialization
initialize()

# Example usage
def run_job_application_system():
    resume_path = ""/content/drive/My Drive/Resume_Template.docx""
    cover_letter_path = ""/content/drive/My Drive/Cover_Letter_Template.docx""
    manager = JobApplicationManager(resume_path, cover_letter_path)
    job_keywords = ""cybersecurity IT project management CISSP PMP""
    locations = [""New York, NY"", ""Remote""]
    jobs_found = manager.search_jobs(job_keywords, locations)
    print(""\nTop matching jobs:"")
    for i, (_, job) in enumerate(jobs_found.head(10).iterrows()):
        print(f""{i+1}. {job['Title']} at {job['Company']} ({job['Location']}) - Match Score: {job['Keywords_Match']}"")
    manager.batch_process_jobs(num_jobs=3)
    print(""\nJob application automation completed!"")

run_job_application_system()"

LINK NUMBER 12
Not enough lines

LINK NUMBER 13

File path: ResnetV50ObjectClassificationHttpTriggerCoPilot/Function1.cs
"   private static readonly InferenceSession session = new InferenceSession(""resnet50.onnx"");

   [FunctionName(""ImageClassification"")]
   public static IActionResult Run(
       [HttpTrigger(AuthorizationLevel.Function, ""post"")] HttpRequest req,
       ILogger log)
   {
      log.LogInformation(""Processing image classification request..."");

      try
      {
         using var ms = new MemoryStream();
         req.Body.CopyTo(ms);
         using var image = Image.FromStream(ms);

         var inputTensor = PreprocessImage(image);

         var inputName = session.InputMetadata.Keys.First();
         var outputName = session.OutputMetadata.Keys.First();
         var result = session.Run(new Dictionary<string, NamedOnnxValue>
            {
                { inputName, NamedOnnxValue.CreateFromTensor(inputName, inputTensor) }
            });

         var predictions = result.First().AsTensor<float>().ToArray();

         return new JsonResult(new { predictions });
      }
      catch (Exception ex)
      {
         log.LogError($""Error: {ex.Message}"");
         return new BadRequestObjectResult(""Invalid image or request."");
      }
   }

   private static Tensor<float> PreprocessImage(Image image)
   {
      var resized = new Bitmap(image, new Size(224, 224));
      var tensorData = new float[1 * 3 * 224 * 224];

      for (int y = 0; y < 224; y++)
      {
         for (int x = 0; x < 224; x++)
         {
            var pixel = resized.GetPixel(x, y);
            tensorData[(0 * 3 * 224 * 224) + (0 * 224 * 224) + (y * 224) + x] = pixel.R / 255.0f;
            tensorData[(0 * 3 * 224 * 224) + (1 * 224 * 224) + (y * 224) + x] = pixel.G / 255.0f;
            tensorData[(0 * 3 * 224 * 224) + (2 * 224 * 224) + (y * 224) + x] = pixel.B / 255.0f;
         }
      }

      return new DenseTensor<float>(tensorData, new[] { 1, 3, 224, 224 });
   }"

LINK NUMBER 14
Not enough lines

LINK NUMBER 15
Not enough lines

LINK NUMBER 16

File path: mappi-api/src/services/fetchTideData.ts
"/**
 * Fetch tide data for a given location.
 *
 * @param locationName - The name of the location to fetch tide data for.
 * @param apiKey - The API key for authenticating with the weather API.
 * @returns A promise that resolves to the tide data for the specified location.
 */
"

LINK NUMBER 17
Not enough lines

LINK NUMBER 18

File path: EightBit/EightBit.UnitTest/DeviceTests.cs
"namespace EightBit.UnitTest
{
    using Microsoft.VisualStudio.TestTools.UnitTesting;
    using EightBit;

    [TestClass]
    public class ChipTests
    {
        [TestMethod]
        public void Bit_ReturnsCorrectBit()
        {
            Assert.AreEqual(0x01, Chip.Bit(0));
            Assert.AreEqual(0x02, Chip.Bit(1));
            Assert.AreEqual(0x80, Chip.Bit(7));
            Assert.AreEqual(0x08, Chip.Bit((byte)3));
        }

        [TestMethod]
        public void SetBit_SetsBitCorrectly()
        {
            Assert.AreEqual(0b00001101, Chip.SetBit(0b00001001, 0b00000100));
            Assert.AreEqual(0b00001101, Chip.SetBit(0b00001101, 0b00000100));
            Assert.AreEqual(0b00001101, Chip.SetBit(0b00001101, 0b00000100, true));
            Assert.AreEqual(0b00001001, Chip.SetBit(0b00001101, 0b00000100, false));
        }

        [TestMethod]
        public void ClearBit_ClearsBitCorrectly()
        {
            Assert.AreEqual(0b00001001, Chip.ClearBit(0b00001101, 0b00000100));
            Assert.AreEqual(0b00001101, Chip.ClearBit(0b00001101, 0b00000100, false));
            Assert.AreEqual(0b00001001, Chip.ClearBit(0b00001101, 0b00000100, true));
        }

        [TestMethod]
        public void HighByte_LowByte_WorkCorrectly()
        {
            ushort value = 0xABCD;
            Assert.AreEqual(0xAB, Chip.HighByte(value));
            Assert.AreEqual(0xCD, Chip.LowByte(value));
            int intValue = 0x1234;
            Assert.AreEqual(0x12, Chip.HighByte(intValue));
            Assert.AreEqual(0x34, Chip.LowByte(intValue));
        }

        [TestMethod]
        public void PromoteByte_DemoteByte_WorkCorrectly()
        {
            Assert.AreEqual(0x3400, Chip.PromoteByte(0x34));
            Assert.AreEqual(0x12, Chip.DemoteByte(0x1234));
        }

        [TestMethod]
        public void HigherPart_LowerPart_WorkCorrectly()
        {
            ushort value = 0xABCD;
            Assert.AreEqual(0xAB00, Chip.HigherPart(value));
            Assert.AreEqual(0xCD, Chip.LowerPart(value));
        }

        [TestMethod]
        public void MakeWord_CreatesCorrectWord()
        {
            Assert.AreEqual(0x1234, Chip.MakeWord(0x34, 0x12));
        }

        [TestMethod]
        public void NibbleMethods_WorkCorrectly()
        {
            byte value = 0xAB;
            Assert.AreEqual(0xA, Chip.HighNibble(value));
            Assert.AreEqual(0xB, Chip.LowNibble(value));
            Assert.AreEqual(0xA0, Chip.HigherNibble(value));
            Assert.AreEqual(0xB, Chip.LowerNibble(value));
            Assert.AreEqual(0xB0, Chip.PromoteNibble(value));
            Assert.AreEqual(0xA, Chip.DemoteNibble(value));
        }

        [TestMethod]
        public void CountBits_ReturnsCorrectCount()
        {
            Assert.AreEqual(0, Chip.CountBits(0));
            Assert.AreEqual(1, Chip.CountBits(1));
            Assert.AreEqual(8, Chip.CountBits(0xFF));
        }

        [TestMethod]
        public void EvenParity_ReturnsCorrectParity()
        {
            Assert.IsTrue(Chip.EvenParity(0)); // 0 bits set
            Assert.IsFalse(Chip.EvenParity(1)); // 1 bit set
            Assert.IsTrue(Chip.EvenParity(3)); // 2 bits set
        }

        [TestMethod]
        public void FindFirstSet_ReturnsCorrectIndex()
        {
            Assert.AreEqual(0, Chip.FindFirstSet(0));
            Assert.AreEqual(1, Chip.FindFirstSet(1));
            Assert.AreEqual(2, Chip.FindFirstSet(2));
            Assert.AreEqual(3, Chip.FindFirstSet(4));
            Assert.AreEqual(5, Chip.FindFirstSet(0b10000));
        }
    }
}"

LINK NUMBER 19
Not enough lines

LINK NUMBER 20
Not enough lines

LINK NUMBER 21
Not enough lines

LINK NUMBER 22

File path: backend/django_emqx/utils.py
"    """"""
    Send a data message via Firebase Cloud Messaging (FCM).

    Args:
        token (str): The recipient's FCM device token.
        msg_id (str): The unique message ID.
        title (str): The title of the message.
        body (str): The body content of the message.

    Returns:
        str: The response from the Firebase messaging service.

    Raises:
        ImportError: If the Firebase Admin SDK is not installed.
    """""""

LINK NUMBER 23
Not enough lines

LINK NUMBER 24
Not enough lines

LINK NUMBER 25
Not enough lines

LINK NUMBER 26
Not enough lines

LINK NUMBER 27
Not enough lines

LINK NUMBER 28

File path: tests/integration_test/test_utils_with_sqlite.py
"        return connection

def test_table_schema_includes_name_and_time(db_setup: sqlite3.Connection):
    """""" Check if the table schema includes 'name' and 'time' columns """"""
    cursor = db_setup.cursor()
    cursor.execute(""PRAGMA table_info(Project)"")
    columns = [column[1] for column in cursor.fetchall()]  # Extract column names
    assert ""name"" in columns, ""'name' column MUST exist in the schema""
    assert ""time"" in columns, ""'time' column MUST exist in the schema""

def test_check_if_project_exists(db_setup: sqlite3.Connection, setup_data: list[tuple]):
    """""" Check if the project exists in the database
     
    This test requires:
    - There is a database we can connect to
    - The database has the required tables created
    - There is data in the tables
    """"""
    for element in setup_data:
        does_exist = check_if_project_exists(element[0])
        assert does_exist, ""This project MUST exist"""

LINK NUMBER 29
Not enough lines

LINK NUMBER 30
Not enough lines

LINK NUMBER 31

File path: packages/api/src/sessions/useGetSessionDetail.ts
"import { CURRENT_GENERATION } from '@depromeet-makers/constant';
import type { UseMutationOptions } from '@tanstack/react-query';
import { useMutation, useQueryClient } from '@tanstack/react-query';

import type { CustomError } from '../base';
import { api } from '../base';
import type { Session } from '../types';

interface EditSessionRequest extends Omit<Session, 'sessionId' | 'generation'> {}

interface EditSessionResponse extends Session {}

const editSession = (sessionId: Session['sessionId'], request: EditSessionRequest) => {
  return api.put<EditSessionResponse>(`/v1/sessions/${sessionId}`, {
    ...request,
    generation: CURRENT_GENERATION,
  });
};

export const useEditSession = (
  sessionId: Session['sessionId'],
  options?: UseMutationOptions<EditSessionResponse, CustomError, EditSessionRequest>,
) => {
  const queryClient = useQueryClient();

  return useMutation({
    mutationFn: (request: EditSessionRequest) => editSession(sessionId, request),
    ...options,
    onSuccess: (...params) => {
      options?.onSuccess?.(...params);

      queryClient.invalidateQueries({
        queryKey: ['sessions'],
      });
    },
  });
};"

LINK NUMBER 32

File path: main.go
"	defer syscall.Close(fd)

	// ソケットをインターフェースにバインド
	addr := syscall.SockaddrLinklayer{
		Protocol: syscall.ETH_P_ALL,
		Ifindex:  iface.Index,
	}
	if err := syscall.Bind(fd, &addr); err != nil {
		log.Fatalf(""Failed to bind raw socket: %v"", err)
	}"

LINK NUMBER 33
Not enough lines

LINK NUMBER 34
Not enough lines

LINK NUMBER 35

File path: src/lib_file.cpp
"/**
 * @brief Checks if the values in a timing arc are monotonically increasing.
 *
 * This function examines the values in the specified timing arc to ensure they are monotonically
 * increasing. It checks monotonicity in two dimensions:
 * 1. Across rows (by output load capacitance)
 * 2. Across columns (by input slew, only if is_slew is true)
 *
 * The function logs detailed information about any non-monotonic values found, including
 * cell name, pin names, and conditional statements (when clause) if present.
 *
 * @param cell The JSON object representing the cell being checked
 * @param pin The JSON object representing the pin being checked
 * @param arc The JSON object representing the timing arc being checked
 * @param timing_arc_name The name of the timing arc to check (e.g., ""cell_rise"", ""cell_fall"")
 * @param is_slew Boolean flag indicating whether to check monotonicity across input slew values
 *
 * @return true if all values in the timing arc are monotonic, false otherwise
 *
 * @throw None, but logs errors or warnings for invalid data formats or non-monotonic values
 */"

LINK NUMBER 36
Not enough lines

LINK NUMBER 37
Error fetching diff

LINK NUMBER 38
Not enough lines

LINK NUMBER 39

File path: apps/admin/src/constants/attendance.ts
"export const ATTENDANCE_STATUS_TEXT_COLOR: Record<ATTENDANCE_STATUS, string> = {
  [ATTENDANCE_STATUS.출석대기]: 'text-gray-300',
  [ATTENDANCE_STATUS.출석]: 'text-green-300',
  [ATTENDANCE_STATUS.지각]: 'text-yellow-300',
  [ATTENDANCE_STATUS.결석]: 'text-red-300',
};
"

LINK NUMBER 40
Not enough lines

LINK NUMBER 41
Not enough lines

LINK NUMBER 42
Not enough lines

LINK NUMBER 43
Not enough lines

LINK NUMBER 44
Not enough lines

LINK NUMBER 45

File path: CopilotCpp8Qeens/CopilotCpp8Qeens.cpp
"// #include <iostream>
// include <vector>
// include <chrono>

import <iostream>;
import <vector>;
import <chrono>;

constexpr int N = 8;
using Board = uint64_t;

constexpr Board col_mask = (1ULL << N) - 1;
constexpr Board diag1_mask = (1ULL << (2 * N - 1)) - 1;
constexpr Board diag2_mask = (1ULL << (2 * N - 1)) - 1;

/*
constexpr Board col_threats[N] = {
    0x0101010101010101ULL,
    0x0202020202020202ULL,
    0x0404040404040404ULL,
    0x0808080808080808ULL,
    0x1010101010101010ULL,
    0x2020202020202020ULL,
    0x4040404040404040ULL,
    0x8080808080808080ULL
};

constexpr Board diag1_threats[2 * N - 1] = {
    0x0000000000000080ULL,
    0x0000000000008040ULL,
    0x0000000000804020ULL,
    0x0000000080402010ULL,
    0x0000008040201008ULL,
    0x0000804020100804ULL,
    0x0080402010080402ULL,
    0x8040201008040201ULL,
    0x4020100804020100ULL,
    0x2010080402010000ULL,
    0x1008040201000000ULL,
    0x0804020100000000ULL,
    0x0402010000000000ULL,
    0x0201000000000000ULL,
    0x0100000000000000ULL
};

constexpr Board diag2_threats[2 * N - 1] = {
    0x0000000000000001ULL,
    0x0000000000000102ULL,
    0x0000000000010204ULL,
    0x0000000001020408ULL,
    0x0000000102040810ULL,
    0x0000010204081020ULL,
    0x0001020408102040ULL,
    0x0102040810204080ULL,
    0x0204081020408000ULL,
    0x0408102040800000ULL,
    0x0810204080000000ULL,
    0x1020408000000000ULL,
    0x2040800000000000ULL,
    0x4080000000000000ULL,
    0x8000000000000000ULL
};
//*/

void solve(Board col, Board diag1, Board diag2, int row, std::vector<Board>& solutions, Board board) {
    if (row == N) [[unlikely]] {
        solutions.push_back(board);
        return;
    }

    Board safe = ~(col | diag1 | diag2) & col_mask;

    while (safe) {
        Board p = safe & (-safe);
        safe -= p;
        solve(col | p, (diag1 | p) << 1, (diag2 | p) >> 1, row + 1, solutions, board | (p << (row * N)));
    }
}

void print_solution(Board board) {
    for (int i = 0; i < N; ++i) {
        for (int j = 0; j < N; ++j) {
            std::cout << ((board & (1ULL << (i * N + j))) ? ""1 "" : ""0 "");
        }
        std::cout << ""\n"";
    }
    std::cout << ""\n"";
}

int main() {
    std::vector<Board> solutions;
    auto start = std::chrono::high_resolution_clock::now();
    solve(0, 0, 0, 0, solutions, 0);
    auto end = std::chrono::high_resolution_clock::now();
    std::chrono::duration<double> elapsed = end - start;

    std::cout << ""Time taken: "" << elapsed.count() << "" seconds\n"";
    std::cout << ""Number of solutions found: "" << solutions.size() << ""\n"";

    for (int i = 0; i < std::min(3, static_cast<int>(solutions.size())); ++i) {
        std::cout << ""Solution "" << i + 1 << "":\n"";
        print_solution(solutions[i]);
    }

    return 0;
}

/*
Time taken: 2.68e-05 seconds
Number of solutions found: 92
Solution 1:
1 0 0 0 0 0 0 0
0 0 0 0 1 0 0 0
0 0 0 0 0 0 0 1
0 0 0 0 0 1 0 0
0 0 1 0 0 0 0 0
0 0 0 0 0 0 1 0
0 1 0 0 0 0 0 0
0 0 0 1 0 0 0 0

Solution 2:
1 0 0 0 0 0 0 0
0 0 0 0 0 1 0 0
0 0 0 0 0 0 0 1
0 0 1 0 0 0 0 0
0 0 0 0 0 0 1 0
0 0 0 1 0 0 0 0
0 1 0 0 0 0 0 0
0 0 0 0 1 0 0 0

Solution 3:
1 0 0 0 0 0 0 0
0 0 0 0 0 0 1 0
0 0 0 1 0 0 0 0
0 0 0 0 0 1 0 0
0 0 0 0 0 0 0 1
0 1 0 0 0 0 0 0
0 0 0 0 1 0 0 0
0 0 1 0 0 0 0 0


E:\Documents and Settings\Pablo\My Documents\My Sources\Cpp8Queens\cpp8Queens\x64\Release\CopilotCpp8Qeens.exe (process 10280) exited with code 0.
Press any key to close this window . . .
*/"

LINK NUMBER 46

File path: QuickQuiz/QuickQuiz/QuizLogic/Model/QuizTests.cs
"using System;
using System.Collections.Generic;
using System.Linq;
using Microsoft.VisualStudio.TestTools.UnitTesting;
using QuickQuiz.QuestionLogic.Model;
using QuickQuiz.QuizLogic.Exceptions;
using QuickQuiz.QuizLogic.Model;

namespace QuickQuiz.Tests
{
    [TestClass]
    public class QuizTests
    {
        [TestMethod]
        public void CreateQuiz_ShouldInitializeQuizWithPlayerName()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>
            {
                new Question(""Question 1""),
                new Question(""Question 2"")
            };

            // Act
            Quiz quiz = Quiz.Create(playerName, questions);

            // Assert
            Assert.AreEqual(playerName, quiz.PlayerName);
            Assert.IsNotNull(quiz.QuizId);
            Assert.IsTrue(quiz.QuestionListReadOnly.Any());
        }

        [TestMethod]
        public void GetNextQuestion_ShouldReturnNextUnansweredQuestion()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>
            {
                new Question(""Question 1""),
                new Question(""Question 2"")
            };
            Quiz quiz = Quiz.Create(playerName, questions);

            // Act
            QuizQuestion nextQuestion = quiz.GetNextQuestion();

            // Assert
            Assert.IsNotNull(nextQuestion);
            Assert.AreEqual(""Question 1"", nextQuestion.OriginalQuestion.Text);
        }

        [TestMethod]
        [ExpectedException(typeof(ActiveQuestionNotAnsweredException))]
        public void GetNextQuestion_ShouldThrowExceptionIfActiveQuestionNotAnswered()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>
            {
                new Question(""Question 1""),
                new Question(""Question 2"")
            };
            Quiz quiz = Quiz.Create(playerName, questions);
            QuizQuestion nextQuestion = quiz.GetNextQuestion();

            // Act
            quiz.GetNextQuestion();
        }

        [TestMethod]
        public void AnswerQuestion_ShouldReturnPlayersAnswer()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>();

            Question question = new Question(""Question 1"");
            questions.Add(question);
            question.AddAnswer(new Answer(""Answer 1"", true));
            question.AddAnswer(new Answer(""Answer 2"", false));

            Quiz quiz = Quiz.Create(playerName, questions);
            QuizQuestion nextQuestion = quiz.GetNextQuestion();

            // Act
            Answer answer = quiz.AnswerQuestion(nextQuestion.QuizQuestionId, nextQuestion.OriginalQuestion.Answers.First().AnswerId);

            // Assert
            Assert.IsNotNull(answer);
            Assert.AreEqual(""Answer 1"", answer.Text);
        }

        [TestMethod]
        [ExpectedException(typeof(Exception), ""There is no Active Question"")]
        public void AnswerQuestion_ShouldThrowExceptionIfNoActiveQuestion()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>();

            Question question = new Question(""Question 1"");
            questions.Add(question);
            question.AddAnswer(new Answer(""Answer 1"", true));
            question.AddAnswer(new Answer(""Answer 2"", false));

            Quiz quiz = Quiz.Create(playerName, questions);

            // Act
            quiz.AnswerQuestion(""invalidQuizQuestionId"", ""invalidAnswerId"");
        }

        [TestMethod]
        [ExpectedException(typeof(Exception), ""You are not answering Active Question"")]
        public void AnswerQuestion_ShouldThrowExceptionIfAnsweringNonActiveQuestion()
        {
            // Arrange
            string playerName = ""Test Player"";
            var questions = new List<Question>();

            Question question = new Question(""Question 1"");
            questions.Add(question);
            question.AddAnswer(new Answer(""Answer 1"", true));
            question.AddAnswer(new Answer(""Answer 2"", false));

            Quiz quiz = Quiz.Create(playerName, questions);
            QuizQuestion nextQuestion = quiz.GetNextQuestion();

            // Act
            quiz.AnswerQuestion(""invalidQuizQuestionId"", ""invalidAnswerId"");
        }
    }
}"

LINK NUMBER 47
Not enough lines

LINK NUMBER 48
Not enough lines

LINK NUMBER 49
Not enough lines

LINK NUMBER 50

File path: vue.config.js
"<template>
  <div>
    <h1>Hello World</h1>
  </div>
</template>

<script>
export default {
  name: 'MainView',
};
</script>

<style scoped>
/* Add any styles specific to MainView here */
</style>"

LINK NUMBER 51

File path: tests/wooODM/product/test_category.py
"import unittest
import unittest
from wooODM.products.category import Category
from wooODM.core import WooCommerce

class TestCategoryModel(unittest.TestCase):

    @classmethod
    def setUpClass(cls):
        # Initialize WooCommerce API with dummy credentials for testing
        WooCommerce.init(
            url="""",
            consumer_key="""",
            consumer_secret=""""
        )
    @classmethod
    def create_test_category(cls, name=""Test Category"", slug=""test-category"", description=""A category for testing""):
        category = Category(
            name=name,
            slug=slug,
            description=description
        )
        return category.save()

    @classmethod
    def delete_test_category(cls, category: Category):
        return category.delete()

    def test_create_category(self):
        category = Category(
            name=""Test Category"",
            slug=""test-category"",
            description=""A category for testing""
        )
        saved_category = category.save()
        self.assertIsNotNone(saved_category.id)
        self.assertEqual(saved_category.name, ""Test Category"")

    def test_get_category(self):
        test_category = TestCategoryModel.create_test_category(name=""Get Test Category"", slug=""get-test-category"")
        category = Category.get(test_category.id)
        self.assertEqual(category.id, test_category.id)
        self.assertEqual(category, test_category)

        test_category.delete()
        with self.assertRaises(Exception):
            Category.get(test_category.id)

    def test_get_all_categories(self):
        categories = Category.all(per_page=5, page=1)
        self.assertIsInstance(categories, list)
        self.assertGreaterEqual(len(categories), 1)
        self.assertIsInstance(categories[0], Category)

    def test_create_incomplete_category(self):
        with self.assertRaises(ValueError):
            Category(
                slug=""incomplete-category""
            ).save()

    def test_create_category_without_slug(self):
        with self.assertRaises(ValueError):
            Category(
                name=""Category without Slug""
            ).save()

    def test_create_category_with_invalid_data(self):
        with self.assertRaises(ValueError):
            Category(
                name=""Invalid Category"",
                slug=""invalid-category"",
                description=123  # Invalid type for description
            ).save()

    def test_update_category_with_invalid_data(self):
        category_id = 1  # Update with a valid category ID
        category = Category.get(category_id)
        category.description = 123  # Invalid type for description
        with self.assertRaises(ValueError):
            category.save()

    def test_smoke(self):
        # Create a category
        category = TestCategoryModel.create_test_category(
            name=""Smoke Test Category"",
            slug=""smoke-test-category"",
            description=""A category for smoke testing""
        )
        self.assertIsNotNone(category.id)
        self.assertEqual(category.name, ""Smoke Test Category"")
        self.assertEqual(category.description, ""A category for smoke testing"")

        # Update the category
        category.name = ""Updated Smoke Test Category""
        updated_category = category.save()
        self.assertEqual(updated_category.name, ""Updated Smoke Test Category"")

        # Delete the category
        delete_response = updated_category.delete()
        self.assertEqual(delete_response, updated_category)

if __name__ == '__main__':
    unittest.main()"

LINK NUMBER 52
Not enough lines

LINK NUMBER 53

File path: routes/stock.js
"const express = require('express');
const axios = require('axios');
const router = express.Router();

// Replace with your Alpha Vantage API key
const API_KEY = 'YOUR_API_KEY';

router.get('/:symbol', async (req, res) => {
  const symbol = req.params.symbol;
  const url = `https://www.alphavantage.co/query?function=TIME_SERIES_INTRADAY&symbol=${symbol}&interval=5min&apikey=${API_KEY}`;

  try {
    const response = await axios.get(url);
    const data = response.data;

    if (data['Error Message']) {
      return res.status(404).render('stock', { error: 'Stock symbol not found' });
    }

    const timeSeries = data['Time Series (5min)'];
    const latestTime = Object.keys(timeSeries)[0];
    const latestData = timeSeries[latestTime];

    const stockInfo = {
      symbol,
      price: latestData['1. open'],
      time: latestTime
    };

    res.render('stock', { stock: stockInfo });
  } catch (error) {
    res.status(500).render('stock', { error: 'Error fetching stock data' });
  }
});

module.exports = router;"

LINK NUMBER 54
Not enough lines

LINK NUMBER 55

File path: humanwrittencode.py
"import time
import psutil
import tracemalloc


def quicksort_human(arr):
    if len(arr) <= 1:
        return arr

    stack = [(0, len(arr) - 1)]
    while stack:
        left, right = stack.pop()
        if left >= right:
            continue

        pivot = arr[right]  # Select last element as pivot
        partition_index = left

        for i in range(left, right):
            if arr[i] < pivot:
                arr[i], arr[partition_index] = arr[partition_index], arr[i]
                partition_index += 1

        arr[partition_index], arr[right] = arr[right], arr[partition_index]

        stack.append((left, partition_index - 1))
        stack.append((partition_index + 1, right))

    return arr


def measure_performance_human():
    arr = [i for i in range(10000, 0, -1)]

    tracemalloc.start()
    start_time = time.time()
    cpu_before = psutil.cpu_percent(interval=None)

    sorted_arr = quicksort_human(arr)  # Human-written sorting

    cpu_after = psutil.cpu_percent(interval=None)
    end_time = time.time()
    mem_usage = tracemalloc.get_traced_memory()

    tracemalloc.stop()

    return {
        ""Execution Time"": end_time - start_time,
        ""Memory Usage"": mem_usage[1] / 1024,  # KB
        ""CPU Usage"": cpu_after - cpu_before
    }


# Run the function
human_results = measure_performance_human()
print(human_results)"

LINK NUMBER 56

File path: src/test/java/com/example/roadsign/StepDefinitions.java
"package com.example.roadsign;

import io.cucumber.java.en.Given;
import io.cucumber.java.en.When;
import io.cucumber.java.en.Then;

public class StepDefinitions {

    @Given(""I have a configured Cucumber project"")
    public void i_have_a_configured_cucumber_project() {
        // Write code here that turns the phrase above into concrete actions
    }

    @When(""I run the tests"")
    public void i_run_the_tests() {
        // Write code here that turns the phrase above into concrete actions
    }

    @Then(""I should see the results"")
    public void i_should_see_the_results() {
        // Write code here that turns the phrase above into concrete actions
    }
}"

LINK NUMBER 57
Not enough lines

LINK NUMBER 58
Not enough lines

LINK NUMBER 59
Not enough lines

LINK NUMBER 60
Not enough lines

LINK NUMBER 61

File path: packages/api/src/sessions/useGetSessionList.ts
"import { CURRENT_GENERATION } from '@depromeet-makers/constant';
import type { UseMutationOptions } from '@tanstack/react-query';
import { useMutation, useQueryClient } from '@tanstack/react-query';

import type { CustomError } from '../base';
import { api } from '../base';
import type { Session } from '../types';

interface CreateSessionRequest extends Omit<Session, 'sessionId' | 'generation'> {}

interface CreateSessionResponse extends Session {}

const createSession = (request: CreateSessionRequest) => {
  return api.post<CreateSessionResponse>('/v1/sessions', { ...request, generation: CURRENT_GENERATION });
};

export const useCreateSession = (
  options?: UseMutationOptions<CreateSessionResponse, CustomError, CreateSessionRequest>,
) => {
  const queryClient = useQueryClient();

  return useMutation({
    mutationFn: createSession,
    ...options,
    onSuccess: (...params) => {
      options?.onSuccess?.(...params);

      queryClient.invalidateQueries({
        queryKey: ['sessions'],
      });
    },
  });
};"

LINK NUMBER 62

File path: packages/api/src/sessions/useGetSessionList.ts
"'use server';

import { cookies } from 'next/headers';
import { COOKIE_KEY } from '@depromeet-makers/constant';
import Cookies from 'js-cookie';

const isClientSide = typeof window !== 'undefined';

/**
 * 인증 토큰을 쿠키와 클라이언트 환경에 저장합니다.
 *
 * @param token 저장할 인증 토큰
 * @returns 저장된 토큰
 */
export const setToken = async (token: string) => {
  const cookieStore = await cookies();

  cookieStore.set(COOKIE_KEY.ACCESS_TOKEN, token);

  if (isClientSide) {
    Cookies.set(COOKIE_KEY.ACCESS_TOKEN, token);
  }

  return token;
};

/**
 * 저장된 인증 토큰을 가져옵니다.
 * 서버 쿠키를 우선 확인하고, 없으면 클라이언트 쿠키를 확인합니다.
 *
 * @returns 저장된 토큰 또는 빈 문자열
 */
export const getToken = async () => {
  const cookieStore = await cookies();
  const serverToken = cookieStore.get(COOKIE_KEY.ACCESS_TOKEN)?.value;

  if (serverToken) return serverToken;

  if (isClientSide) {
    const clientToken = Cookies.get(COOKIE_KEY.ACCESS_TOKEN);

    if (clientToken) {
      setToken(clientToken);
      return clientToken;
    }
  }

  return '';
};

/**
 * 저장된 인증 토큰 제거
 */
export const removeToken = async () => {
  const cookieStore = await cookies();

  cookieStore.delete(COOKIE_KEY.ACCESS_TOKEN);

  if (isClientSide) {
    Cookies.remove(COOKIE_KEY.ACCESS_TOKEN);
  }
};"

LINK NUMBER 63
Not enough lines

LINK NUMBER 64
Not enough lines

LINK NUMBER 65
Not enough lines

LINK NUMBER 66

File path: app/__init__.py
"# 🚜 Smart Fleet Maintenance API

## 📌 Overview
The **Smart Fleet Maintenance API** is designed to help track maintenance schedules, predict equipment failures, and log service records for agricultural machinery.

## 🔥 Features
- **Register Equipment** – Store machine details (model, purchase date, usage hours).
- **Maintenance Scheduler** – Predict maintenance needs based on usage.
- **Service Logging** – Track service records and update machine status.
- **Alerts & Notifications** – Notify users when maintenance is due.

## 🏗️ Tech Stack
- **Backend:** FastAPI (Python)
- **Database:** SQLite / PostgreSQL
- **Testing:** Pytest
- **Data Validation:** Pydantic
- **CI/CD:** GitHub Actions (for automated testing)

## 🏁 Getting Started

### **🔹 Prerequisites**
Ensure you have the following installed:
- Python 3.9+
- pip
- Virtual environment tool (venv or conda)

### **🔹 Installation**
```bash
# Clone the repository
git clone https://github.com/yourusername/smart-fleet-maintenance-api.git
cd smart-fleet-maintenance-api

# Set up a virtual environment
python -m venv venv
source venv/bin/activate  # On Windows: venv\Scripts\activate

# Install dependencies
pip install -r requirements.txt
```

### **🔹 Running the API**
```bash
uvicorn main:app --reload
```
API will be available at: `http://127.0.0.1:8000`

### **🔹 Running Tests**
```bash
pytest tests/
```

## 📂 Project Structure
```
smart-fleet-maintenance-api/
│── main.py          # Entry point for FastAPI application
│── database.py      # Database connection setup
│── requirements.txt # List of dependencies
|── models/
    |── models.py      # Pydantic models for request validation
│── tests/           # Pytest test cases
│── routers/         # API route handlers
│   ├── equipment.py
│   ├── maintenance.py
│── .github/workflows/ci.yml  # GitHub Actions for CI/CD
│── README.md        # Project documentation
```

## 🚀 API Endpoints
### 1️⃣ Register Equipment
**POST** `/equipment/`
#### Request Body:
```json
{
  ""name"": ""Tractor X"",
  ""model"": ""TX-500"",
  ""purchase_date"": ""2023-01-15"",
  ""usage_hours"": 100
}
```
#### Response:
```json
{
  ""id"": ""UUID"",
  ""name"": ""Tractor X"",
  ""model"": ""TX-500"",
  ""purchase_date"": ""2023-01-15"",
  ""usage_hours"": 100
}
```

### 2️⃣ Get Equipment List
**GET** `/equipment/`

More endpoints will be added as features are implemented.

## 📌 Future Enhancements
- **Machine Learning for Predictive Maintenance**
- **Fleet Analytics Dashboard (Power BI)**
- **Real-Time Alerts via WebSockets**

---
### 💡 Contributing
Feel free to open an issue or submit a pull request.

### 📜 License
MIT License. See `LICENSE` for details.
"

LINK NUMBER 67

File path: apps/admin/src/app/(admin)/session/(data)/session.ts
"import type { Dispatch, PropsWithChildren, SetStateAction } from 'react';
import { createContext, useContext, useState } from 'react';

export interface MarkerType {
  id: string;
  position: {
    lat: number;
    lng: number;
  };
  placeName: string;
  addressName: string;
}

interface UsersContextType {
  markers: MarkerType[];
  setMarkers: Dispatch<SetStateAction<MarkerType[]>>;
  selectedPlace?: MarkerType;
  setSelectedPlace: Dispatch<SetStateAction<MarkerType | undefined>>;
}

const KaKaoMapContext = createContext<UsersContextType | null>(null);

const KaKaoMapProvider = ({ children }: PropsWithChildren) => {
  const [markers, setMarkers] = useState<MarkerType[]>([]);
  const [selectedPlace, setSelectedPlace] = useState<MarkerType>();

  return (
    <KaKaoMapContext.Provider value={{ markers, setMarkers, selectedPlace, setSelectedPlace }}>
      {children}
    </KaKaoMapContext.Provider>
  );
};

export const useKaKaoMap = () => {
  const kaKaoMapContext = useContext(KaKaoMapContext);

  if (!kaKaoMapContext) {
    throw new Error('<kaKaoMapContext /> 내부에서 useKaKaoMap을 사용할 수 있어요.');
  }

  return kaKaoMapContext;
};

export default KaKaoMapProvider;"

LINK NUMBER 68
Not enough lines

LINK NUMBER 69
Not enough lines

LINK NUMBER 70
Not enough lines

LINK NUMBER 71

File path: packages/api/src/sessions/useDeleteSession.ts
"import type { UseMutationOptions } from '@tanstack/react-query';
import { useMutation, useQueryClient } from '@tanstack/react-query';

import type { CustomError } from '../base';
import { api } from '../base';
import type { Session } from '../types';

interface DeleteSessionRequest extends Pick<Session, 'sessionId'> {}

interface DeleteSessionResponse extends Session {}

const deleteSession = ({ sessionId }: DeleteSessionRequest) => {
  return api.delete<DeleteSessionResponse>(`/v1/sessions/${sessionId}`);
};

export const useDeleteSession = (
  options?: UseMutationOptions<DeleteSessionResponse, CustomError, DeleteSessionRequest>,
) => {
  const queryClient = useQueryClient();

  return useMutation({
    mutationFn: ({ sessionId }: DeleteSessionRequest) => deleteSession({ sessionId }),
    ...options,
    onSuccess: (...params) => {
      options?.onSuccess?.(...params);

      queryClient.invalidateQueries({ queryKey: ['sessions'] });
    },
  });
};"

LINK NUMBER 72

File path: turingmachine.py
"for i in range(0,len(tape)):
	try:
		tape[i] = int(tape[i])
	except ValueError:
		try:
			assert tape[i] in symbols
		except AssertionError:
			tape[i] = ""A""
	else:
		try:
			assert tape[i] in symbols
		except AssertionError:
			tape[i] = 0"

LINK NUMBER 73
Not enough lines

LINK NUMBER 74
Not enough lines

LINK NUMBER 75

File path: python/mandelbrot_set/mandelbrot.py
"from mpi4py import MPI
import numpy as np
import matplotlib.pyplot as plt

def mandelbrot(c, max_iter):
    z = c
    for n in range(max_iter):
        if abs(z) > 2:
            return n
        z = z*z + c
    return max_iter

def mandelbrot_set(xmin, xmax, ymin, ymax, width, height, max_iter):
    r1 = np.linspace(xmin, xmax, width)
    r2 = np.linspace(ymin, ymax, height)
    n3 = np.empty((width, height), dtype=int)
    for i in range(width):
        for j in range(height):
            n3[i, j] = mandelbrot(r1[i] + 1j*r2[j], max_iter)
    return n3

def main():
    comm = MPI.COMM_WORLD
    rank = comm.Get_rank()
    size = comm.Get_size()

    # Parameters for the Mandelbrot set
    xmin, xmax, ymin, ymax = -2.0, 1.0, -1.5, 1.5
    width, height = 800, 800
    max_iter = 256

    # Divide the work among processes
    local_width = width // size
    local_xmin = xmin + rank * local_width * (xmax - xmin) / width
    local_xmax = xmin + (rank + 1) * local_width * (xmax - xmin) / width

    # Each process computes its part of the Mandelbrot set
    local_mandelbrot = mandelbrot_set(local_xmin, local_xmax, ymin, ymax, local_width, height, max_iter)

    # Gather the results to the root process
    if rank == 0:
        mandelbrot_image = np.empty((width, height), dtype=int)
    else:
        mandelbrot_image = None

    comm.Gather(local_mandelbrot, mandelbrot_image, root=0)

    # Root process visualizes the result
    if rank == 0:
        plt.imshow(mandelbrot_image.T, extent=[xmin, xmax, ymin, ymax], cmap='hot')
        plt.colorbar()
        plt.title(""Mandelbrot Set"")
        plt.xlabel(""Re"")
        plt.ylabel(""Im"")
        plt.savefig(""mandelbrot_set.png"")
        # plt.show()

if __name__ == ""__main__"":
    main()"

LINK NUMBER 76

File path: main.py
"import torch
from torchvision import datasets, transforms

import torch.nn as nn
import torch.optim as optim

# Define the neural network model
class MNISTModel(nn.Module):
    def __init__(self):
        super(MNISTModel, self).__init__()
        self.layer1 = nn.Sequential(
            nn.Conv2d(1, 32, kernel_size=5, stride=1, padding=2),
            nn.ReLU(),
            nn.MaxPool2d(kernel_size=2, stride=2))
        self.layer2 = nn.Sequential(
            nn.Conv2d(32, 64, kernel_size=5, stride=1, padding=2),
            nn.ReLU(),
            nn.MaxPool2d(kernel_size=2, stride=2))
        self.fc1 = nn.Linear(7*7*64, 1000)
        self.fc2 = nn.Linear(1000, 10)

    def forward(self, x):
        out = self.layer1(x)
        out = self.layer2(out)
        out = out.view(out.size(0), -1)
        out = self.fc1(out)
        out = self.fc2(out)
        return out

# Load the MNIST dataset
transform = transforms.Compose([
    transforms.ToTensor(),
    transforms.Normalize((0.1307,), (0.3081,))
])

print(""before train"")
train_dataset = datasets.MNIST(root='./data', train=True, transform=transform, download=True)
print(""before test"")
test_dataset = datasets.MNIST(root='./data', train=False, transform=transform)
print(""data set complete"")

train_loader = torch.utils.data.DataLoader(dataset=train_dataset, batch_size=64, shuffle=True)
test_loader = torch.utils.data.DataLoader(dataset=test_dataset, batch_size=64, shuffle=False)

# Initialize the model, loss function, and optimizer
model = MNISTModel()
criterion = nn.CrossEntropyLoss()
optimizer = optim.Adam(model.parameters(), lr=0.001)

# Train the model
num_epochs = 5
for epoch in range(num_epochs):
    for i, (images, labels) in enumerate(train_loader):
        outputs = model(images)
        loss = criterion(outputs, labels)
        
        optimizer.zero_grad()
        loss.backward()
        optimizer.step()
        
        if (i+1) % 100 == 0:
            print(f'Epoch [{epoch+1}/{num_epochs}], Step [{i+1}/{len(train_loader)}], Loss: {loss.item():.4f}')

# Test the model
model.eval()
with torch.no_grad():
    correct = 0
    total = 0
    for images, labels in test_loader:
        outputs = model(images)
        _, predicted = torch.max(outputs.data, 1)
        total += labels.size(0)
        correct += (predicted == labels).sum().item()

    print(f'Accuracy of the model on the 10000 test images: {100 * correct / total} %')"

LINK NUMBER 77

File path: packages/api/src/sessions/useGetSessionList.ts
"'use client';

import * as React from 'react';
import * as RadioGroupPrimitive from '@radix-ui/react-radio-group';
import { Circle } from 'lucide-react';

import { cn } from '@/lib/utils';

const RadioGroup = React.forwardRef<
  React.ElementRef<typeof RadioGroupPrimitive.Root>,
  React.ComponentPropsWithoutRef<typeof RadioGroupPrimitive.Root>
>(({ className, ...props }, ref) => {
  return <RadioGroupPrimitive.Root className={cn('grid gap-2', className)} {...props} ref={ref} />;
});
RadioGroup.displayName = RadioGroupPrimitive.Root.displayName;

const RadioGroupItem = React.forwardRef<
  React.ElementRef<typeof RadioGroupPrimitive.Item>,
  React.ComponentPropsWithoutRef<typeof RadioGroupPrimitive.Item>
>(({ className, ...props }, ref) => {
  return (
    <RadioGroupPrimitive.Item
      ref={ref}
      className={cn(
        'aspect-square h-4 w-4 rounded-full border border-primary text-primary ring-offset-background focus:outline-none focus-visible:ring-2 focus-visible:ring-ring focus-visible:ring-offset-2 disabled:cursor-not-allowed disabled:opacity-50',
        className,
      )}
      {...props}
    >
      <RadioGroupPrimitive.Indicator className=""flex items-center justify-center"">
        <Circle className=""h-2.5 w-2.5 fill-current text-current"" />
      </RadioGroupPrimitive.Indicator>
    </RadioGroupPrimitive.Item>
  );
});
RadioGroupItem.displayName = RadioGroupPrimitive.Item.displayName;

export { RadioGroup, RadioGroupItem };"

LINK NUMBER 78

File path: src/classes/edge-of-the-empire-dice/proficiency-die.ts
"	/**
	 * Maps the value of the proficiency die to the corresponding array of `EdgeOfTheEmpireDiceSymbol` results.
	 *
	 * @returns {EdgeOfTheEmpireDiceSymbol[]} An array of `EdgeOfTheEmpireDiceSymbol` representing the result of the die roll.
	 *
	 * The mapping is as follows:
	 * - 1: Blank
	 * - 2, 3: Success
	 * - 4, 5: Success, Success
	 * - 6: Advantage
	 * - 7, 8, 9: Success, Advantage
	 * - 10, 11: Advantage, Advantage
	 * - 12: Triumph
	 * - Default: Empty array
	 */"

LINK NUMBER 79
Not enough lines

LINK NUMBER 80

File path: packages/api/src/base/token.ts
"import type { UseMutationOptions } from '@tanstack/react-query';
import { useMutation } from '@tanstack/react-query';

import type { CustomError } from '../base';
import { api } from '../base';
import { setAccessToken, setRefreshToken } from '../base/token';

interface PostAuthTestResponse {
  accessToken: string;
  refreshToken: string;
}

const postAuthTest = () => {
  return api.post<PostAuthTestResponse>('/v1/auth/test');
};

export const useAuthTest = (options?: UseMutationOptions<PostAuthTestResponse, CustomError>) =>
  useMutation({
    mutationFn: postAuthTest,
    ...options,
    onSuccess: async (data, ...params) => {
      await setAccessToken(data.accessToken);
      await setRefreshToken(data.refreshToken);

      options?.onSuccess?.(data, ...params);
    },
  });"

LINK NUMBER 81
Not enough lines

LINK NUMBER 82
Not enough lines

LINK NUMBER 83
Not enough lines

LINK NUMBER 84

File path: 1873E.cpp
"#include<bits/stdc++.h>
using namespace std;
#define ll long long

ll solve(){
    ll n,x;
    cin>>n>>x;

    vector<ll> a(n);
    for(int i=0;i<n;i++){
        cin>>a[i];
    }

    ll l=1;
    ll r=INT_MAX;
    ll ans=0;

    while(l<=r){
        ll mid=(l+r)/2;
        ll water=0;
        for(int i=0;i<n;i++){
            if(mid>a[i]){
                water+=(mid-a[i]);
            }
        }

        if(water<=x){
            ans=max(ans,mid);
            l=mid+1;
        }
        else{
            r=mid-1;
        }
    }

    return ans;

}

int main(){
    int t;
    cin>>t;
    while(t--){
        cout<<solve()<<endl;
    }
}"

LINK NUMBER 85
Not enough lines

LINK NUMBER 86
Not enough lines

LINK NUMBER 87

File path: java/BonusCalculator.java
"/**
 * Encapsulates the logic for calculating bonuses based on gross salary.
 */
public class BonusCalculator {
    // Constants
    private static final double BONUS_RATE = 0.10;

    /**
     * Calculates the bonus based on the given gross salary.
     *
     * @param grossSalary the gross salary of the employee
     * @return the calculated bonus
     */
    public double calculateBonus(double grossSalary) {
        return grossSalary * BONUS_RATE;
    }
}"

LINK NUMBER 88

File path: java/PayrollSystem.java
"import java.util.ArrayList;
import java.util.Collections;
import java.util.Comparator;
import java.util.List;

/**
 * Manages the overall payroll system, including initializing employees, sorting employees, calculating net salaries,
 * calculating department totals, and displaying information.
 */
public class PayrollSystem {
    // Constants
    private static final int MAX_EMPLOYEES = 5;
    private static final double TAX_RATE = 0.20;
    private static final double BONUS_RATE = 0.10;
    private static final double DEDUCTION_RATE = 0.05;

    // Attributes
    private List<Employee> employees = new ArrayList<>();
    private List<DepartmentTotal> departmentTotals = new ArrayList<>();

    /**
     * Initializes the employee data with hardcoded values.
     */
    public void initializeEmployees() {
        employees.add(new Employee(""E001"", ""Alice Johnson"", ""HR"", 70000.00));
        employees.add(new Employee(""E002"", ""Bob Smith"", ""IT"", 85000.00));
        employees.add(new Employee(""E003"", ""Charlie Brown"", ""Finance"", 60000.00));
        employees.add(new Employee(""E004"", ""David Wilson"", ""IT"", 95000.00));
        employees.add(new Employee(""E005"", ""Eve Davis"", ""HR"", 75000.00));
    }

    /**
     * Sorts the employees by their IDs.
     */
    public void sortEmployees() {
        Collections.sort(employees, Comparator.comparing(Employee::getId));
    }

    /**
     * Calculates the net salaries for all employees.
     */
    public void calculateNetSalaries() {
        BonusCalculator bonusCalculator = new BonusCalculator();
        for (Employee employee : employees) {
            double bonus = bonusCalculator.calculateBonus(employee.getGrossSalary());
            double deductions = employee.getGrossSalary() * DEDUCTION_RATE;
            double taxDeduction = employee.getGrossSalary() * TAX_RATE;
            double netSalary = employee.getGrossSalary() + bonus - taxDeduction - deductions;

            employee.setBonus(bonus);
            employee.setDeductions(deductions);
            employee.setTaxDeduction(taxDeduction);
            employee.setNetSalary(netSalary);
        }
    }

    /**
     * Calculates the total salaries for each department.
     */
    public void calculateDepartmentTotals() {
        for (Employee employee : employees) {
            DepartmentTotal departmentTotal = departmentTotals.stream()
                .filter(dt -> dt.getDepartmentName().equals(employee.getDepartment()))
                .findFirst()
                .orElseGet(() -> {
                    DepartmentTotal newDeptTotal = new DepartmentTotal(employee.getDepartment());
                    departmentTotals.add(newDeptTotal);
                    return newDeptTotal;
                });

            departmentTotal.setTotalSalary(departmentTotal.getTotalSalary() + employee.getNetSalary());
        }
    }

    /**
     * Displays the employee payroll information.
     */
    public void displayEmployees() {
        System.out.println(""Employee Payroll Information"");
        System.out.println(""-----------------------------"");
        for (Employee employee : employees) {
            System.out.println(""Employee ID: "" + employee.getId());
            System.out.println(""Name: "" + employee.getName());
            System.out.println(""Department: "" + employee.getDepartment());
            System.out.println(""Gross Salary: $"" + employee.getGrossSalary());
            System.out.println(""Bonus: $"" + employee.getBonus());
            System.out.println(""Deductions: $"" + employee.getDeductions());
            System.out.println(""Tax Deduction: $"" + employee.getTaxDeduction());
            System.out.println(""Net Salary: $"" + employee.getNetSalary());
            System.out.println(""-----------------------------"");
        }
    }

    /**
     * Displays the department salary totals.
     */
    public void displayDepartmentTotals() {
        System.out.println(""Department Salary Totals"");
        System.out.println(""-----------------------------"");
        for (DepartmentTotal departmentTotal : departmentTotals) {
            System.out.println(""Department: "" + departmentTotal.getDepartmentName());
            System.out.println(""Total Salary: $"" + departmentTotal.getTotalSalary());
            System.out.println(""-----------------------------"");
        }
    }

    /**
     * The main method to run the payroll system.
     *
     * @param args command-line arguments
     */
    public static void main(String[] args) {
        PayrollSystem payrollSystem = new PayrollSystem();
        payrollSystem.initializeEmployees();
        payrollSystem.sortEmployees();
        payrollSystem.calculateNetSalaries();
        payrollSystem.calculateDepartmentTotals();
        payrollSystem.displayEmployees();
        payrollSystem.displayDepartmentTotals();
    }
}"

LINK NUMBER 89

File path: java/DepartmentTotal.java
"/**
 * Represents the total salary for a department with attributes like department name and total salary.
 */
public class DepartmentTotal {
    // Attributes
    private String departmentName;
    private double totalSalary;

    /**
     * Initializes the department with the given name and sets the total salary to 0.
     *
     * @param departmentName the name of the department
     */
    public DepartmentTotal(String departmentName) {
        this.departmentName = departmentName;
        this.totalSalary = 0;
    }

    // Getters and setters for all attributes

    public String getDepartmentName() {
        return departmentName;
    }

    public void setDepartmentName(String departmentName) {
        this.departmentName = departmentName;
    }

    public double getTotalSalary() {
        return totalSalary;
    }

    public void setTotalSalary(double totalSalary) {
        this.totalSalary = totalSalary;
    }
}"

LINK NUMBER 90
Not enough lines

LINK NUMBER 91
Not enough lines

LINK NUMBER 92

File path: multiple-file/java/Employee.java
"package java;

/**
 * Represents an employee with attributes like ID, name, department, gross salary, bonus, deductions, net salary, and tax deduction.
 */
public class Employee {
    // Attributes
    private String id;
    private String name;
    private String department;
    private double grossSalary;
    private double bonus;
    private double deductions;
    private double netSalary;
    private double taxDeduction;

    /**
     * Initializes the employee with the given ID, name, department, and gross salary.
     * Other attributes are initialized to default values.
     *
     * @param id          the employee's ID
     * @param name        the employee's name
     * @param department  the department the employee belongs to
     * @param grossSalary the employee's gross salary
     */
    public Employee(String id, String name, String department, double grossSalary) {
        this.id = id;
        this.name = name;
        this.department = department;
        this.grossSalary = grossSalary;
        this.bonus = 0.0;
        this.deductions = 0.0;
        this.netSalary = 0.0;
        this.taxDeduction = 0.0;
    }

    // Getters and setters for all attributes

    public String getId() {
        return id;
    }

    public void setId(String id) {
        this.id = id;
    }

    public String getName() {
        return name;
    }

    public void setName(String name) {
        this.name = name;
    }

    public String getDepartment() {
        return department;
    }

    public void setDepartment(String department) {
        this.department = department;
    }

    public double getGrossSalary() {
        return grossSalary;
    }

    public void setGrossSalary(double grossSalary) {
        this.grossSalary = grossSalary;
    }

    public double getBonus() {
        return bonus;
    }

    public void setBonus(double bonus) {
        this.bonus = bonus;
    }

    public double getDeductions() {
        return deductions;
    }

    public void setDeductions(double deductions) {
        this.deductions = deductions;
    }

    public double getNetSalary() {
        return netSalary;
    }

    public void setNetSalary(double netSalary) {
        this.netSalary = netSalary;
    }

    public double getTaxDeduction() {
        return taxDeduction;
    }

    public void setTaxDeduction(double taxDeduction) {
        this.taxDeduction = taxDeduction;
    }
}"

LINK NUMBER 93

File path: src/test/java/sa/com/cloudsolutions/antikythera/generator/UnitTestGeneratorTest.java
"

class VariableInitializationModifierTest {

    @Test
    void shouldModifySimpleVariableInitialization() {
        String code = """"""
            public void testMethod() {
                String test = ""old"";
                int other = 5;
            }
            """""";
        MethodDeclaration method = StaticJavaParser.parseMethodDeclaration(code);
        StringLiteralExpr newValue = new StringLiteralExpr(""new"");

        VariableInitializationModifier modifier = new VariableInitializationModifier(""test"", newValue);
        MethodDeclaration result = (MethodDeclaration) modifier.visit(method, null);

        assertTrue(result.toString().contains(""String test = \""new\""""));
        assertTrue(result.toString().contains(""int other = 5""));
    }

    @Test
    void shouldNotModifyWhenVariableNotFound() {
        String code = """"""
            public void testMethod() {
                String existingVar = ""old"";
            }
            """""";
        MethodDeclaration method = StaticJavaParser.parseMethodDeclaration(code);
        IntegerLiteralExpr newValue = new IntegerLiteralExpr(""42"");

        VariableInitializationModifier modifier = new VariableInitializationModifier(""nonexistentVar"", newValue);
        MethodDeclaration result = (MethodDeclaration) modifier.visit(method, null);

        assertEquals(method.toString(), result.toString());
    }

    @Test
    void shouldModifyFirstOccurrenceOnly() {
        String code = """"""
            public void testMethod() {
                int target = 1;
                String other = ""middle"";
                int target = 3;
            }
            """""";
        MethodDeclaration method = StaticJavaParser.parseMethodDeclaration(code);
        IntegerLiteralExpr newValue = new IntegerLiteralExpr(""42"");

        VariableInitializationModifier modifier = new VariableInitializationModifier(""target"", newValue);
        MethodDeclaration result = (MethodDeclaration) modifier.visit(method, null);

        String modifiedCode = result.toString();
        assertTrue(modifiedCode.contains(""int target = 42""));
        assertTrue(modifiedCode.contains(""int target = 3""));
        assertEquals(1, modifiedCode.split(""42"").length - 1);
    }
}"

LINK NUMBER 94
Not enough lines

LINK NUMBER 95
Not enough lines

LINK NUMBER 96
Not enough lines

LINK NUMBER 97
Not enough lines

LINK NUMBER 98

File path: neon_grid.py
"import pygame
import random
from player import Player
from enemy import Enemy
from data_shard import DataShard
from platform import Platform
from neon_grid import NeonGrid

# Initialize Pygame
pygame.init()

# Screen dimensions
SCREEN_WIDTH = 800
SCREEN_HEIGHT = 600

# Colors
BLACK = (0, 0, 0)
WHITE = (255, 255, 255)

# Create the screen
screen = pygame.display.set_mode((SCREEN_WIDTH, SCREEN_HEIGHT))
pygame.display.set_caption(""Neon Malfunction"")

# Clock to control the frame rate
clock = pygame.time.Clock()

# Create game objects
player = Player(SCREEN_WIDTH // 2, SCREEN_HEIGHT // 2)
enemies = [Enemy(random.randint(0, SCREEN_WIDTH), random.randint(0, SCREEN_HEIGHT)) for _ in range(5)]
data_shards = [DataShard(random.randint(0, SCREEN_WIDTH), random.randint(0, SCREEN_HEIGHT)) for _ in range(10)]
platforms = [Platform(random.randint(0, SCREEN_WIDTH), random.randint(0, SCREEN_HEIGHT)) for _ in range(5)]
neon_grid = NeonGrid()

# Game loop
running = True
while running:
    for event in pygame.event.get():
        if event.type == pygame.QUIT:
            running = False

    # Player controls
    keys = pygame.key.get_pressed()
    if keys[pygame.K_SPACE]:
        player.flap()
    if keys[pygame.K_a]:
        player.move_left()
    if keys[pygame.K_d]:
        player.move_right()

    # Update game objects
    player.update()
    for enemy in enemies:
        enemy.update()
    for data_shard in data_shards:
        data_shard.update()
    for platform in platforms:
        platform.update()
    neon_grid.update()

    # Collision detection
    for enemy in enemies:
        if player.collides_with(enemy):
            if player.position.y < enemy.position.y:
                enemies.remove(enemy)
            else:
                running = False

    for data_shard in data_shards:
        if player.collides_with(data_shard):
            player.score += data_shard.value
            data_shards.remove(data_shard)

    # Draw everything
    screen.fill(BLACK)
    player.draw(screen)
    for enemy in enemies:
        enemy.draw(screen)
    for data_shard in data_shards:
        data_shard.draw(screen)
    for platform in platforms:
        platform.draw(screen)
    neon_grid.draw(screen)

    # Update the display
    pygame.display.flip()

    # Cap the frame rate
    clock.tick(60)

# Quit Pygame
pygame.quit()"

LINK NUMBER 99

File path: src/static/main.js
"// This file contains the JavaScript code for client-side interactions in the todo tracker app.

document.addEventListener('DOMContentLoaded', function() {
    const addTodoForm = document.getElementById('add-todo-form');
    const todoList = document.getElementById('todo-list');

    addTodoForm.addEventListener('submit', function(event) {
        event.preventDefault();
        const title = document.getElementById('todo-title').value;
        const body = document.getElementById('todo-body').value;
        const dueDate = document.getElementById('todo-due-date').value;

        // Add new todo item
        fetch('/todos', {
            method: 'POST',
            headers: {
                'Content-Type': 'application/json'
            },
            body: JSON.stringify({ title, body, due_date: dueDate })
        })
        .then(response => response.json())
        .then(data => {
            // Append new todo item to the list
            const todoItem = document.createElement('li');
            todoItem.textContent = `${data.title} - ${data.body}`;
            todoList.appendChild(todoItem);
            addTodoForm.reset();
        });
    });

    // Function to filter todos by completion status
    document.getElementById('filter-completed').addEventListener('change', function() {
        const showCompleted = this.checked;
        const todos = todoList.getElementsByTagName('li');
        for (let todo of todos) {
            if (showCompleted && !todo.classList.contains('completed')) {
                todo.style.display = 'none';
            } else {
                todo.style.display = 'list-item';
            }
        }
    });

    // Function to sort todos by creation date or due date
    document.getElementById('sort-todos').addEventListener('change', function() {
        const sortBy = this.value;
        const todosArray = Array.from(todoList.getElementsByTagName('li'));
        todosArray.sort((a, b) => {
            const aDate = new Date(a.dataset.creationTime);
            const bDate = new Date(b.dataset.creationTime);
            return sortBy === 'due_date' ? aDate - bDate : bDate - aDate;
        });
        todoList.innerHTML = '';
        todosArray.forEach(todo => todoList.appendChild(todo));
    });
});"

LINK NUMBER 100
Not enough lines

LINK NUMBER 101

File path: go-base64/go-base64.go
"package main

import (
    ""encoding/base64""
    ""fmt""
    ""io/ioutil""
    ""os""
)

// EncodeFile encodes the content of the input file and writes it to the output file
func EncodeFile(inputFile, outputFile string) error {
    data, err := ioutil.ReadFile(inputFile)
    if err != nil {
        return fmt.Errorf(""failed to read file: %w"", err)
    }

    encodedData := base64.StdEncoding.EncodeToString(data)

    err = ioutil.WriteFile(outputFile, []byte(encodedData), 0644)
    if err != nil {
        return fmt.Errorf(""failed to write file: %w"", err)
    }

    return nil
}

// DecodeFile decodes the base64 content of the input file and writes it to the output file
func DecodeFile(inputFile, outputFile string) error {
    data, err := ioutil.ReadFile(inputFile)
    if err != nil {
        return fmt.Errorf(""failed to read file: %w"", err)
    }

    decodedData, err := base64.StdEncoding.DecodeString(string(data))
    if err != nil {
        return fmt.Errorf(""failed to decode base64 data: %w"", err)
    }

    err = ioutil.WriteFile(outputFile, decodedData, 0644)
    if err != nil {
        return fmt.Errorf(""failed to write file: %w"", err)
    }

    return nil
}

func ShowUsage() {
    fmt.Println(""Usage:"")
    fmt.Println(""  go-base64 encode <input file> <output file>"")
    fmt.Println(""  go-base64 decode <input file> <output file>"")
}

func main() {
    if len(os.Args) < 4 {
        ShowUsage()
        return
    }

    command := os.Args[1]
    inputFile := os.Args[2]
    outputFile := os.Args[3]

    var err error
    switch command {
    case ""encode"":
        err = EncodeFile(inputFile, outputFile)
    case ""decode"":
        err = DecodeFile(inputFile, outputFile)
    default:
        ShowUsage()
        return
    }

    if err != nil {
        fmt.Printf(""Error: %v\n"", err)
    } else {
        fmt.Printf(""Success: %s completed\n"", command)
    }
}"

LINK NUMBER 102
Not enough lines

LINK NUMBER 103
Not enough lines

LINK NUMBER 104
Not enough lines

LINK NUMBER 105

File path: main.go
"package main

import (
	""log""
	""net""

	""golang.org/x/net/ethernet""
)

func main() {
	// UDPの受信ポートを指定
	listenAddr := "":12345""
	conn, err := net.ListenPacket(""udp"", listenAddr)
	if err != nil {
		log.Fatalf(""Failed to listen on UDP port: %v"", err)
	}
	defer conn.Close()

	log.Printf(""Listening for UDP packets on %s"", listenAddr)

	// 転送先のMACアドレスを指定
	targetMAC := net.HardwareAddr{0x00, 0x11, 0x22, 0x33, 0x44, 0x55} // 例: 00:11:22:33:44:55
	interfaceName := ""eth0""                                           // 使用するネットワークインターフェース名

	// ネットワークインターフェースを取得
	iface, err := net.InterfaceByName(interfaceName)
	if err != nil {
		log.Fatalf(""Failed to get interface %s: %v"", interfaceName, err)
	}

	// RAWソケットを作成
	rawConn, err := ethernet.ListenPacket(iface)
	if err != nil {
		log.Fatalf(""Failed to create raw socket: %v"", err)
	}
	defer rawConn.Close()

	buffer := make([]byte, 1500) // 最大パケットサイズ

	for {
		// UDPパケットを受信
		n, addr, err := conn.ReadFrom(buffer)
		if err != nil {
			log.Printf(""Error reading UDP packet: %v"", err)
			continue
		}

		log.Printf(""Received %d bytes from %s"", n, addr)

		// Ethernetフレームを構築
		ethFrame := &ethernet.Frame{
			Destination: targetMAC,
			Source:      iface.HardwareAddr,
			EtherType:   ethernet.EtherTypeIPv4,
			Payload:     buffer[:n],
		}

		// Ethernetフレームをシリアライズ
		ethData, err := ethFrame.MarshalBinary()
		if err != nil {
			log.Printf(""Error serializing Ethernet frame: %v"", err)
			continue
		}

		// RAWソケットで送信
		_, err = rawConn.WriteTo(ethData, &ethernet.Addr{HardwareAddr: targetMAC})
		if err != nil {
			log.Printf(""Error sending Ethernet frame: %v"", err)
			continue
		}

		log.Printf(""Forwarded packet to MAC %s"", targetMAC)
	}
}"

LINK NUMBER 106

File path: bimbink-web-app/src/firebase.js
"import { initializeApp } from ""firebase/app"";
import { getFirestore } from ""firebase/firestore"";
import { getAuth } from ""firebase/auth"";
import { getAnalytics } from ""firebase/analytics"";

const firebaseConfig = {
    apiKey: ""YOUR_API_KEY"",
    authDomain: ""YOUR_PROJECT_ID.firebaseapp.com"",
    projectId: ""YOUR_PROJECT_ID"",
    storageBucket: ""YOUR_PROJECT_ID.appspot.com"",
    messagingSenderId: ""YOUR_MESSAGING_SENDER_ID"",
    appId: ""YOUR_APP_ID""
};

const app = initializeApp(firebaseConfig);
const db = getFirestore(app);
const auth = getAuth(app);
const analytics = getAnalytics(app);

export { db, auth };"

LINK NUMBER 107

File path: app/src/models/commit.ts
"  /**
   * Whether or not the message was generated by Copilot
   * (optional, default: false)
   */
  readonly messageGeneratedByCopilot?: boolean"

LINK NUMBER 108
Not enough lines

LINK NUMBER 109
Not enough lines

LINK NUMBER 110
Not enough lines

LINK NUMBER 111

File path: app/src/models/commit.ts
"  /**
   * Whether or not the message was generated by Copilot
   * (optional, default: false)
   */
  readonly messageGeneratedByCopilot?: boolean"

LINK NUMBER 112
Not enough lines

LINK NUMBER 113
Not enough lines

LINK NUMBER 114
Not enough lines

LINK NUMBER 115

File path: test_sqlite/populate_database.py
"import json
import sqlite3
import os
from tqdm import tqdm  # For progress bars

def populate_test_database(json_file, db_file):
    """"""Populate SQLite database from the repeat domain JSON file""""""
    print(f""Creating database from {json_file}..."")
    
    # Connect to SQLite DB
    conn = sqlite3.connect(db_file)
    cursor = conn.cursor()
    
    # Create tables using schema file
    # Get the directory of the script to find the schema file
    script_dir = os.path.dirname(os.path.abspath(__file__))
    schema_file_path = os.path.join(script_dir, ""database_schema.sql"")
    
    with open(schema_file_path, ""r"") as schema_file:
        schema = schema_file.read()
        conn.executescript(schema)
    
    # Load JSON data
    with open(json_file, ""r"") as f:
        data = json.load(f)
    
    # Track processed IDs to avoid duplicates
    processed_genes = {}
    processed_proteins = {}
    processed_transcripts = {}
    processed_exons = {}

    # Process each repeat entry
    for repeat in tqdm(data, desc=""Processing repeats""):
        if not isinstance(repeat, dict) or not repeat:
            continue  # Skip empty entries
        
        # Get gene info
        gene_name = repeat.get(""geneName"", """")
        if not gene_name:
            continue
        
        # Insert gene if not already processed
        if gene_name not in processed_genes:
            aliases = repeat.get(""aliases"", """")
            if isinstance(aliases, list):
                aliases = "","".join(aliases)
            
            cursor.execute(""""""
                INSERT INTO genes (gene_name, aliases, chromosome, location)
                VALUES (?, ?, ?, ?)
            """""", (
                gene_name,
                aliases,
                repeat.get(""chrom"", """"),
                f""{repeat.get('chrom', '')}:{repeat.get('chromStart', '')}_{repeat.get('chromEnd', '')}""
            ))
            processed_genes[gene_name] = cursor.lastrowid
        
        gene_id = processed_genes[gene_name]
        
        # Process protein
        uniprot_id = repeat.get(""uniProtId"", """")
        if uniprot_id and uniprot_id not in processed_proteins:
            status = repeat.get(""status"", """")
            
            # Extract length from position if possible: ""amino acids 343-389 on protein Q6TDP4""
            position = repeat.get(""position"", """")
            length = 0
            if position and isinstance(position, str):
                try:
                    parts = position.split()
                    if len(parts) >= 3:
                        pos = parts[2].split(""-"")
                        if len(pos) == 2:
                            length = int(pos[1]) - int(pos[0]) + 1
                except:
                    pass
            
            cursor.execute(""""""
                INSERT INTO proteins (protein_id, gene_id, length, description, status)
                VALUES (?, ?, ?, ?, ?)
            """""", (
                uniprot_id,
                gene_id,
                length,
                f""Protein for {gene_name}"",
                status
            ))
            processed_proteins[uniprot_id] = uniprot_id
        
        # Process repeat domain
        amino_start = None
        amino_end = None
        if repeat.get(""position"") and isinstance(repeat.get(""position""), str):
            position = repeat.get(""position"")
            try:
                # Extract positions from ""amino acids 343-389 on protein Q6TDP4""
                parts = position.split()
                if len(parts) >= 3:
                    pos = parts[2].split(""-"")
                    if len(pos) == 2:
                        amino_start = int(pos[0])
                        amino_end = int(pos[1])
            except:
                pass
        
        # Calculate sequence length from blockSizes
        sequence_length = 0
        block_sizes = repeat.get(""blockSizes"", [])
        if isinstance(block_sizes, list):
            for size in block_sizes:
                try:
                    sequence_length += int(size)
                except:
                    pass
        
        # Insert repeat
        cursor.execute(""""""
            INSERT INTO repeats (protein_id, repeat_type, start_pos, end_pos, sequence)
            VALUES (?, ?, ?, ?, ?)
        """""", (
            uniprot_id,
            repeat.get(""repeatType"", """"),
            amino_start,
            amino_end,
            ""N"" * sequence_length  # Placeholder sequence of Ns
        ))
        repeat_id = cursor.lastrowid
        
        # Process exon information if available
        if ""ensembl_exon_info"" not in repeat:
            continue
            
        exon_info = repeat.get(""ensembl_exon_info"", {})
        if not exon_info or ""transcripts"" not in exon_info:
            continue
            
        # Process each transcript
        for transcript_data in exon_info.get(""transcripts"", []):
            transcript_id = transcript_data.get(""transcript_id"")
            if not transcript_id:
                continue
                
            # Insert transcript if not already processed
            if transcript_id not in processed_transcripts:
                cursor.execute(""""""
                    INSERT INTO transcripts (transcript_id, gene_id, description)
                    VALUES (?, ?, ?)
                """""", (
                    transcript_id,
                    gene_id,
                    f""{transcript_data.get('transcript_name', '')} ({transcript_data.get('biotype', '')})""
                ))
                processed_transcripts[transcript_id] = transcript_id
            
            # Create repeat_transcript relationship
            genomic_start = repeat.get(""chromStart"", 0)
            genomic_end = repeat.get(""chromEnd"", 0)
            
            # Convert exon mapping to JSON string
            exon_mapping = json.dumps([
                {
                    ""exon_id"": exon.get(""exon_id"", """"),
                    ""exon_number"": exon.get(""exon_number"", 0),
                    ""overlap_bp"": exon.get(""overlap_bp"", 0),
                    ""overlap_percentage"": exon.get(""overlap_percentage"", 0),
                    ""coding_percentage"": exon.get(""coding_percentage"", 0)
                } for exon in transcript_data.get(""containing_exons"", [])
            ])
            
            cursor.execute(""""""
                INSERT INTO repeat_transcripts (
                    repeat_id, transcript_id, genomic_start, genomic_end, exon_mapping
                ) VALUES (?, ?, ?, ?, ?)
            """""", (
                repeat_id, 
                transcript_id, 
                genomic_start, 
                genomic_end, 
                exon_mapping
            ))
            
            # Process exons in this transcript
            for exon_data in transcript_data.get(""containing_exons"", []):
                exon_id = exon_data.get(""exon_id"")
                if not exon_id or exon_id in processed_exons:
                    continue
                    
                # Estimate exon size from overlap percentage
                exon_size = 0
                if exon_data.get(""overlap_percentage"") and exon_data.get(""overlap_bp""):
                    try:
                        exon_size = int(exon_data.get(""overlap_bp"") * 100 / exon_data.get(""overlap_percentage""))
                    except:
                        pass
                
                # Calculate if skipping would preserve reading frame
                frame_preserving = exon_size % 3 == 0
                
                cursor.execute(""""""
                    INSERT INTO exons (
                        exon_id, gene_id, length, frame_preserving
                    ) VALUES (?, ?, ?, ?)
                """""", (
                    exon_id,
                    gene_id,
                    exon_size,
                    frame_preserving
                ))
                processed_exons[exon_id] = exon_id
                
                # Create transcript_exon relationship
                cursor.execute(""""""
                    INSERT INTO transcript_exons (
                        transcript_id, exon_id, exon_number
                    ) VALUES (?, ?, ?)
                """""", (
                    transcript_id,
                    exon_id,
                    exon_data.get(""exon_number"", 0)
                ))
                
                # Create repeat_exon relationship
                cursor.execute(""""""
                    INSERT INTO repeat_exons (
                        repeat_id, exon_id, overlap_bp, overlap_percentage
                    ) VALUES (?, ?, ?, ?)
                """""", (
                    repeat_id,
                    exon_id,
                    exon_data.get(""overlap_bp"", 0),
                    exon_data.get(""overlap_percentage"", 0)
                ))
    
    # Commit all changes
    conn.commit()
    
    # Print some statistics
    cursor.execute(""SELECT COUNT(*) FROM genes"")
    gene_count = cursor.fetchone()[0]
    
    cursor.execute(""SELECT COUNT(*) FROM proteins"")
    protein_count = cursor.fetchone()[0]
    
    cursor.execute(""SELECT COUNT(*) FROM repeats"")
    repeat_count = cursor.fetchone()[0]
    
    cursor.execute(""SELECT COUNT(*) FROM exons"")
    exon_count = cursor.fetchone()[0]
    
    print(f""\nDatabase populated successfully:"")
    print(f""  - {gene_count} genes"")
    print(f""  - {protein_count} proteins"")
    print(f""  - {repeat_count} repeat domains"")
    print(f""  - {exon_count} exons"")
    
    conn.close()
    
    print(f""\nDatabase created at: {db_file}"")

if __name__ == ""__main__"":
    # Define file paths - simplified to use the same directory as the script
    script_dir = os.path.dirname(os.path.abspath(__file__))
    
    json_file = os.path.join(script_dir, ""1000_test_exons_hg38_repeats.json"")
    db_file = os.path.join(script_dir, ""tandem_repeats.db"")
    
    # Create the database
    populate_test_database(json_file, db_file)"

LINK NUMBER 116
Not enough lines

LINK NUMBER 117

File path: 2024/Common/common.c
"ssize_t getline(char **lineptr, size_t *n, FILE *stream) {
    if (lineptr == NULL || n == NULL || stream == NULL) {
        return -1;
    }

    char *buf = *lineptr;
    size_t size = *n;
    int c = 0;
    size_t len = 0;

    if (buf == NULL || size == 0) {
        size = 128;
        buf = (char *)malloc(size);
        if (buf == NULL) {
            return -1;
        }
    }

    while ((c = fgetc(stream)) != EOF) {
        if (len + 1 >= size) {
            size *= 2;
            char *new_buf = (char *)realloc(buf, size);
            if (new_buf == NULL) {
                free(buf);
                return -1;
            }
            buf = new_buf;
        }
        buf[len++] = (char)c;
        if (c == '\n') {
            break;
        }
    }

    if (len == 0 && c == EOF) {
        return -1;
    }

    buf[len] = '\0';
    *lineptr = buf;
    *n = size;

    return len;
}
"

LINK NUMBER 118
Not enough lines

LINK NUMBER 119
Not enough lines

LINK NUMBER 120

File path: scripts/checkAllowance.ts
"/** 
 * Use environment variables to provide arguments 
 * export contractAddress=""""
 * export spender=""""
 * export amount=""""
 * Run the script with hardhat run 
 * npx hardhat run scripts/approveTokens.ts --network sepolia 
 */
import { ethers } from ""hardhat"";

async function approveTokens(contractAddress: string, spender: string, amount: string) {
  const GLDToken = await ethers.getContractFactory(""GLDToken"");
  const token = GLDToken.attach(contractAddress);

  const tx = await token.approve(spender, ethers.utils.parseUnits(amount, 18));
  console.log(`Approved ${amount} GLD for ${spender}. Transaction hash: ${tx.hash}`);
}

const contractAddress = process.env.contractAddress as string;
const spender = process.env.spender as string;
const amount = process.env.amount as string;

approveTokens(contractAddress, spender, amount)
  .then(() => process.exit(0))
  .catch((error) => {
    console.error(error);
    process.exit(1);
  });"

LINK NUMBER 121
Not enough lines

LINK NUMBER 122
Not enough lines

LINK NUMBER 123
Not enough lines

LINK NUMBER 124
Not enough lines

LINK NUMBER 125

File path: src/com/sandeep/java8/clockutils/Java8date.java
"/**
 * The Java8date class demonstrates various ways to work with dates using the java.time package introduced in Java 8.
 * It includes examples of getting the current date, creating specific dates, handling time zones, and working with epoch days.
 * 
 * <p>Examples include:</p>
 * <ul>
 *   <li>Getting the current date</li>
 *   <li>Creating a LocalDate with specific year, month, and day</li>
 *   <li>Handling invalid date inputs</li>
 *   <li>Getting the current date in a specific time zone</li>
 *   <li>Creating a date from the epoch day</li>
 *   <li>Creating a date from the day of the year</li>
 * </ul>
 * 
 * <p>Note: Some examples are commented out to demonstrate potential exceptions.</p>
 * 
 * @author Sandeep
 */"

LINK NUMBER 126

File path: handler_polka.go
"package main

import (
	""encoding/json""
	""fmt""
	""log""
	""net/http""

	""github.com/google/uuid""
)

// Handlers for Polka payment processing webhooks

// NOTE: Ideally negative responses should include a retry-after header. Not including for now.
func (cfg *apiConfig) handlePolkaWebhook(response http.ResponseWriter, request *http.Request) {
	// Parse request params
	var params struct {
		Event string `json:""event""`
		Data struct {
			UserID uuid.UUID `json:""user_id""`
		}
	}

	decoder := json.NewDecoder(request.Body)
	err := decoder.Decode(&params)
	if err != nil {
		msg := fmt.Sprintf(""polka: Error decoding webhook params: %s\n"", err)
		log.Println(msg)
		respondWithError(response, http.StatusBadRequest, msg)
		return
	}

	switch params.Event {
	case ""user.upgraded"":
		cfg.handleUserUpgrade(response, request, params.Data.UserID)
	default:
		msg := fmt.Sprintf(""polka: Unknown event type: %s\n"", params.Event)
		log.Println(msg)
		respondWithError(response, http.StatusNoContent, msg)
		return
	}
}"

LINK NUMBER 127
Not enough lines

LINK NUMBER 128
Not enough lines

LINK NUMBER 129

File path: src/index.ts
"{
  ""name"": ""@zereight/mcp-confluence"",
  ""version"": ""1.0.1"",
  ""description"": ""MCP server for using the Confluence API"",
  ""license"": ""MIT"",
  ""author"": ""zereight"",
  ""type"": ""module"",
  ""private"": false,
  ""bin"": ""./build/index.js"",
  ""files"": [
    ""build""
  ],
  ""publishConfig"": {
    ""access"": ""public""
  },
  ""engines"": {
    ""node"": "">=14""
  },
  ""scripts"": {
    ""build"": ""tsc && node -e \""require('fs').chmodSync('build/index.js', '755')\"""",
    ""prepare"": ""npm run build"",
    ""watch"": ""tsc --watch"",
    ""inspector"": ""npx @modelcontextprotocol/inspector build/index.js"",
    ""start"": ""node build/index.js""
  },
  ""dependencies"": {
    ""@modelcontextprotocol/sdk"": ""0.6.0"",
    ""axios"": ""^1.7.9"",
    ""mcp-framework"": ""^0.1.12"",
    ""okhttp"": ""^1.1.0""
  },
  ""devDependencies"": {
    ""@types/node"": ""^20.11.24"",
    ""typescript"": ""^5.7.2""
  }
}"

LINK NUMBER 130
Not enough lines

LINK NUMBER 131
Not enough lines

LINK NUMBER 132
Not enough lines

LINK NUMBER 133
Not enough lines

LINK NUMBER 134

File path: tests/utils/date.test.ts
"import { describe, it, expect } from ""vitest"";

import { SmartIterator, ValueException } from ""../../src/index.js"";
import { Curve } from ""../../src/index.js"";

describe(""Curve"", () =>
{
    describe(""Linear"", () =>
    {
        it(""Should return an instance of `SmartIterator`"", () =>
        {
            const iterator = Curve.Linear(5);

            expect(iterator).toBeInstanceOf(SmartIterator);
        });
        it(""Should generate a linear sequence of values"", () =>
        {
            const values = Array.from(Curve.Linear(5));

            expect(values).toEqual([0, 0.25, 0.5, 0.75, 1]);
        });
    });

    describe(""Exponential"", () =>
    {
        it(""Should return an instance of `SmartIterator`"", () =>
        {
            const iterator = Curve.Exponential(6);

            expect(iterator).toBeInstanceOf(SmartIterator);
        });

        it(""Should generate an exponential sequence of values with default base"", () =>
        {
            const values = Array.from(Curve.Exponential(6));

            expect(values).toEqual([0, 0.04000000000000001, 0.16000000000000003, 0.36, 0.6400000000000001, 1]);
        });
        it(""Should generate an exponential sequence of values with custom base"", () =>
        {
            const values = Array.from(Curve.Exponential(6, 3));

            expect(values).toEqual(
                [0, Math.pow(1 / 5, 3), Math.pow(2 / 5, 3), Math.pow(3 / 5, 3), Math.pow(4 / 5, 3), 1]
            );
        });

        it(""Should throw a `ValueException` if base is negative"", () =>
        {
            expect(() => Curve.Exponential(6, -1)).toThrow(ValueException);
        });
    });
});"

LINK NUMBER 135
Too many lines

LINK NUMBER 136

File path: src/classes/edge-of-the-empire-dice/proficiency-die.ts
"	/**
	 * Maps the value of the proficiency die to the corresponding array of `EdgeOfTheEmpireDiceSymbol` results.
	 *
	 * @returns {EdgeOfTheEmpireDiceSymbol[]} An array of `EdgeOfTheEmpireDiceSymbol` representing the result of the die roll.
	 *
	 * The mapping is as follows:
	 * - 1: Blank
	 * - 2, 3: Success
	 * - 4, 5: Success, Success
	 * - 6: Advantage
	 * - 7, 8, 9: Success, Advantage
	 * - 10, 11: Advantage, Advantage
	 * - 12: Triumph
	 * - Default: Empty array
	 */"

LINK NUMBER 137

File path: js/timers.js
"document.getElementById('createTimers').addEventListener('click', function() {
    const numTimers = parseInt(document.getElementById('numTimers').value);
    const defaultTime = parseInt(document.getElementById('defaultTime').value);
    const timersContainer = document.getElementById('timersContainer');
    timersContainer.innerHTML = '';

    for (let i = 1; i <= numTimers; i++) {
        const timerDiv = document.createElement('div');
        timerDiv.className = 'timer';
        timerDiv.innerHTML = `
            <input type=""text"" value=""${i}"">
            <div class=""time"">${defaultTime}</div>
            <div class=""buttons"">
                <button class=""start"">Start</button>
                <button class=""pause"">Pause</button>
                <button class=""reset"">Reset</button>
            </div>
        `;
        timersContainer.appendChild(timerDiv);

        const timeDisplay = timerDiv.querySelector('.time');
        let timeLeft = defaultTime;
        let interval;

        timerDiv.querySelector('.start').addEventListener('click', function() {
            if (interval) return;
            interval = setInterval(() => {
                if (timeLeft > 0) {
                    timeLeft--;
                    timeDisplay.textContent = timeLeft;
                    if (timeLeft <= 10) {
                        timeDisplay.classList.add('warning');
                    }
                } else {
                    clearInterval(interval);
                    timeDisplay.textContent = 'Times Up!';
                }
            }, 1000);
        });

        timerDiv.querySelector('.pause').addEventListener('click', function() {
            clearInterval(interval);
            interval = null;
        });

        timerDiv.querySelector('.reset').addEventListener('click', function() {
            clearInterval(interval);
            interval = null;
            timeLeft = defaultTime;
            timeDisplay.textContent = timeLeft;
            timeDisplay.classList.remove('warning');
        });
    }
});"

LINK NUMBER 138
Not enough lines

LINK NUMBER 139
Not enough lines

LINK NUMBER 140
Not enough lines

LINK NUMBER 141

File path: src/common/middleware/auth.middleware.ts
"
    req.user = {
      userId: decoded.userId,
      role: decoded.role,
    };

    next();
  } catch (error) {
    logger.error(""JWT verification failed"", { error });
    return next(new UnauthorizedError(""Invalid or expired token""));
  }"

LINK NUMBER 142
Not enough lines

LINK NUMBER 143

File path: hilbert12.py
"import numpy as np

def generate_hilbert_matrix(size):
    """"""Generates a Hilbert matrix of given size.""""""
    hilbert_matrix = np.array([[1 / (i + j + 1) for j in range(size)] for i in range(size)])
    return hilbert_matrix

if __name__ == ""__main__"":
    size = 12
    hilbert_matrix = generate_hilbert_matrix(size)
    print(hilbert_matrix)"

LINK NUMBER 144
Not enough lines

LINK NUMBER 145

File path: kalman.py
"import numpy as np

class KalmanFilter:
    def __init__(self, dt, u, std_acc, x_std_meas, y_std_meas, z_std_meas):
        # Define sampling time
        self.dt = dt

        # Define the 3D state variables (position and velocity)
        self.u = u
        self.x = np.zeros((6, 1))

        # Define the state transition matrix
        self.A = np.array([[1, 0, 0, self.dt, 0, 0],
                           [0, 1, 0, 0, self.dt, 0],
                           [0, 0, 1, 0, 0, self.dt],
                           [0, 0, 0, 1, 0, 0],
                           [0, 0, 0, 0, 1, 0],
                           [0, 0, 0, 0, 0, 1]])

        # Define the control input matrix
        self.B = np.array([[0.5 * self.dt**2, 0, 0],
                           [0, 0.5 * self.dt**2, 0],
                           [0, 0, 0.5 * self.dt**2],
                           [self.dt, 0, 0],
                           [0, self.dt, 0],
                           [0, 0, self.dt]])

        # Define the measurement mapping matrix
        self.H = np.array([[1, 0, 0, 0, 0, 0],
                           [0, 1, 0, 0, 0, 0],
                           [0, 0, 1, 0, 0, 0]])

        # Initial covariance matrix
        self.P = np.eye(6)

        # Define the process noise covariance matrix
        self.Q = std_acc**2 * np.array([[0.25 * self.dt**4, 0, 0, 0.5 * self.dt**3, 0, 0],
                                        [0, 0.25 * self.dt**4, 0, 0, 0.5 * self.dt**3, 0],
                                        [0, 0, 0.25 * self.dt**4, 0, 0, 0.5 * self.dt**3],
                                        [0.5 * self.dt**3, 0, 0, self.dt**2, 0, 0],
                                        [0, 0.5 * self.dt**3, 0, 0, self.dt**2, 0],
                                        [0, 0, 0.5 * self.dt**3, 0, 0, self.dt**2]])

        # Define the measurement noise covariance matrix
        self.R = np.array([[x_std_meas**2, 0, 0],
                           [0, y_std_meas**2, 0],
                           [0, 0, z_std_meas**2]])

    def predict(self):
        # Predict the state
        self.x = np.dot(self.A, self.x) + np.dot(self.B, self.u)

        # Predict the error covariance
        self.P = np.dot(np.dot(self.A, self.P), self.A.T) + self.Q

        return self.x

    def update(self, z):
        # Compute the Kalman Gain
        S = np.dot(self.H, np.dot(self.P, self.H.T)) + self.R
        K = np.dot(np.dot(self.P, self.H.T), np.linalg.inv(S))

        # Update the state estimate
        self.x = self.x + np.dot(K, (z - np.dot(self.H, self.x)))

        # Update the error covariance
        self.P = self.P - np.dot(np.dot(K, self.H), self.P)

        return self.x

# Example usage
dt = 0.1
u = np.array([[0], [0], [0]])  # No control input
std_acc = 0.1
x_std_meas = 0.1
y_std_meas = 0.1
z_std_meas = 0.1

kf = KalmanFilter(dt, u, std_acc, x_std_meas, y_std_meas, z_std_meas)

measurements = [np.array([[1], [2], [3]]), np.array([[1.1], [2.1], [3.1]]), np.array([[0.9], [1.9], [2.9]])]

for z in measurements:
    kf.predict()
    state = kf.update(z)
    print(""Updated state:\n"", state)"

LINK NUMBER 146

File path: hilbert.py
"import numpy as np

def hilbert_matrix(n):
    """"""Generate an n x n Hilbert matrix.""""""
    H = np.zeros((n, n))
    for i in range(n):
        for j in range(n):
            H[i, j] = 1 / (i + j + 1)
    return H

# Generate a 10x10 Hilbert matrix
n = 10
H = hilbert_matrix(n)

# Print the matrix
print(""10x10 Hilbert Matrix:"")
print(H)"

LINK NUMBER 147

File path: matrix.py
"import numpy as np

def parse_matrix_input():
    rows = []
    print(""Enter matrix rows one by one. Type 'end' to finish:"")
    while True:
        row = input()
        if row.lower() == 'end':
            break
        rows.append([float(num) for num in row.split()])
    return np.array(rows)

def matrix_calculator():
    while True:
        print(""\nMatrix Calculator - Supported operations: add, sub, mul, div, exit"")
        operation = input(""Enter operation: "").strip().lower()

        if operation == 'exit':
            break

        print(""Enter the first matrix:"")
        matrix1 = parse_matrix_input()

        print(""Enter the second matrix:"")
        matrix2 = parse_matrix_input()

        if operation == 'add':
            try:
                result = np.add(matrix1, matrix2)
                print(""Result:\n"", result)
            except ValueError as e:
                print(""Error:"", e)
        elif operation == 'sub':
            try:
                result = np.subtract(matrix1, matrix2)
                print(""Result:\n"", result)
            except ValueError as e:
                print(""Error:"", e)
        elif operation == 'mul':
            try:
                result = np.matmul(matrix1, matrix2)
                print(""Result:\n"", result)
            except ValueError as e:
                print(""Error:"", e)
        elif operation == 'div':
            try:
                result = np.divide(matrix1, matrix2)
                print(""Result:\n"", result)
            except ValueError as e:
                print(""Error:"", e)
        else:
            print(""Invalid operation"")

if __name__ == ""__main__"":
    matrix_calculator()"

LINK NUMBER 148

File path: calc2.py
"import math

def text_based_calculator():
    while True:
        print(""\nSelect operation:"")
        print(""1. Addition (+)"")
        print(""2. Subtraction (-)"")
        print(""3. Multiplication (*)"")
        print(""4. Division (/)"")
        print(""5. Trigonometry (sin, cos, tan)"")
        print(""6. Square Root (sqrt)"")
        print(""7. Square (sq)"")
        print(""8. Exponential (exp)"")
        print(""9. Exit"")

        choice = input(""Enter choice: "")

        if choice == '9':
            break

        if choice in ['1', '2', '3', '4']:
            num1 = float(input(""Enter first number: ""))
            num2 = float(input(""Enter second number: ""))

            if choice == '1':
                print(""Result:"", num1 + num2)
            elif choice == '2':
                print(""Result:"", num1 - num2)
            elif choice == '3':
                print(""Result:"", num1 * num2)
            elif choice == '4':
                if num2 != 0:
                    print(""Result:"", num1 / num2)
                else:
                    print(""Error: Division by zero"")

        elif choice in ['5']:
            trig_operation = input(""Enter trigonometric function (sin, cos, tan): "")
            angle = float(input(""Enter angle in radians: ""))
            if trig_operation == 'sin':
                print(""Result:"", math.sin(angle))
            elif trig_operation == 'cos':
                print(""Result:"", math.cos(angle))
            elif trig_operation == 'tan':
                print(""Result:"", math.tan(angle))
            else:
                print(""Invalid trigonometric function"")

        elif choice == '6':
            num = float(input(""Enter number: ""))
            print(""Result:"", math.sqrt(num))

        elif choice == '7':
            num = float(input(""Enter number: ""))
            print(""Result:"", num * num)

        elif choice == '8':
            num = float(input(""Enter number: ""))
            print(""Result:"", math.exp(num))

        else:
            print(""Invalid choice"")

if __name__ == ""__main__"":
    text_based_calculator()"

LINK NUMBER 149
Not enough lines

LINK NUMBER 150
Not enough lines

LINK NUMBER 151

File path: static/music.js
"        ['e2', 8], ['e2', 8], ['p', 8], ['e2', 8], ['p', 8], ['c2', 8], ['e2', 8],
        ['g2', 4], ['p', 4], ['g1', 4], ['p', 4],
        ['c2', 4], ['p', 8], ['g1', 8], ['g1', 8], ['p', 8], ['e1', 4],
        ['p', 8], ['a1', 4], ['h1', 8], ['h1', 8], ['ais1', 8], ['ais1', 4],
        ['g1', 6], ['e2', 6], ['g2', 6], ['a2', 4], ['f2', 8], ['g2', 8],
        ['p', 8], ['e2', 4], ['c2', 8], ['d2', 8], ['h1', 4], ['p', 8],
        ['c2', 4], ['p', 8], ['g1', 8], ['g1', 8], ['p', 8], ['e1', 4],
        ['p', 4], ['a1', 4], ['h1', 8], ['h1', 8], ['ais1', 8], ['ais1', 4],
        ['g1', 6], ['e2', 6], ['g2', 6], ['a2', 4], ['f2', 8], ['g2', 8],
        ['p', 8], ['e2', 4], ['c2', 8], ['d2', 8], ['h1', 4], ['p', 8],
        ['p', 4], ['g2', 8], ['fis2', 8], ['fis2', 8], ['dis2', 4], ['e2', 8],
        ['p', 8], ['gis1', 8], ['a1', 8], ['c2', 8], ['p', 8], ['a1', 8], ['c2', 8], ['d2', 8],
        ['p', 4], ['g2', 8], ['fis2', 8], ['fis2', 8], ['dis2', 4], ['e2', 8],
        ['p', 8], ['c3', 4], ['c3', 8], ['c3', 4], ['p', 4],
        ['p', 4], ['g2', 8], ['fis2', 8], ['fis2', 8], ['dis2', 4], ['e2', 8],
        ['p', 8], ['gis1', 8], ['a1', 8], ['c2', 8], ['p', 8], ['a1', 8], ['c2', 8], ['d2', 8],"

LINK NUMBER 152

File path: flask-web-app/static/js/script.js
"// script.js

document.addEventListener('DOMContentLoaded', function() {
    const states = [
        {
            state: ""Uttar Pradesh"",
            language: ""Hindi"",
            message: ""शुभ दीपावली (Hindi, Uttar Pradesh)"",
            imageUrl: ""https://diwaliimages.blob.core.windows.net/images/Uttarpradesh.jpg""
        },
        {
            state: ""Punjab"",
            language: ""Punjabi"",
            message: ""ਦੀਵਾਲੀ ਮੁਬਾਰਕ (Punjabi, Punjab)"",
            imageUrl: ""https://diwaliimages.blob.core.windows.net/images/Punjab.jpg""
        },
        {
            state: ""Gujarat"",
            language: ""Gujarati"",
            message: ""શુભ દિવાળી (Gujarati, Gujarat)"",
            imageUrl: ""https://diwaliimages.blob.core.windows.net/images/Gujrat.jpg""
        },
        {
            state: ""Assam"",
            language: ""Assamese"",
            message: ""শুভ দীপাবলি (Assamese, Assam)"",
            imageUrl: ""https://diwaliimages.blob.core.windows.net/images/Assam.jpg""
        },
        {
            state: ""Tamil Nadu"",
            language: ""Tamil"",
            message: ""இனிய தீபாவளி (Tamil, Tamil Nadu)"",
            imageUrl: ""https://diwaliimages.blob.core.windows.net/images/Tamilnadu.jpg""
        }
    ];

    let currentIndex = 0;

    function updateMessageAndImage() {
        currentIndex = (currentIndex + 1) % states.length;
        const stateInfo = states[currentIndex];
        document.getElementById('message').textContent = stateInfo.message;
        document.getElementById('image').src = stateInfo.imageUrl;
    }

    document.addEventListener('click', function() {
        updateMessageAndImage();
    });
});"

LINK NUMBER 153

File path: backend/src/sdk.cpp
"
void SDK::handleAddStation(const nlohmann::json& json)
{
    std::optional<std::string> callsign;
    std::optional<int> frequency;

    if (json[""value""].contains(""callsign"")) {
        callsign = json[""value""][""callsign""];
    }
    if (json[""value""].contains(""frequency"")) {
        frequency = json[""value""][""frequency""];
    }

    if (callsign.has_value() && frequency.has_value()) {
        PLOG_ERROR << ""Both callsign and frequency specified. Only one should be specified."";
        return;
    }

    if (!callsign.has_value() && !frequency.has_value()) {
        PLOG_ERROR << ""Neither callsign nor frequency specified. One must be specified."";
        return;
    }

    auto allRadios = mClient->getRadioState();
    for (const auto& [freq, state] : allRadios) {
        if ((callsign.has_value() && state.stationName == callsign.value())
            || (frequency.has_value() && freq == frequency.value())) {
            this->publishStationState(
                this->buildStationStateJson(state.stationName, static_cast<int>(freq)));
            return;
        }
    }

    if (callsign.has_value()) {
        mClient->GetStation(callsign.value());
    } else if (frequency.has_value()) {
        mClient->AddFrequency(frequency.value(), """");
    }
}"

LINK NUMBER 154
Too many lines

LINK NUMBER 155
Not enough lines

LINK NUMBER 156
Not enough lines

LINK NUMBER 157
Not enough lines

LINK NUMBER 158
Not enough lines

LINK NUMBER 159
Not enough lines

LINK NUMBER 160
Not enough lines

LINK NUMBER 161
Not enough lines

LINK NUMBER 162
Not enough lines

LINK NUMBER 163
Not enough lines

LINK NUMBER 164
Not enough lines

LINK NUMBER 165
Not enough lines

LINK NUMBER 166
Not enough lines

LINK NUMBER 167

File path: TestArena/Blog/Common/NavigationUtils/SiteMap.cs
"        new(""Understanding the Single Responsibility Principle"",
            ""/blog/software-practices-solid-srp"",
            new DateTime(2025, 5, 10),
            ""images/blog/software-practices/solid-srp/banner.png"",
            [""Software Practices"", ""SOLID"", ""SRP""], false),"

LINK NUMBER 168
Not enough lines

LINK NUMBER 169

File path: src/features/follow/index.ts
"import { NextFunction, Request, Response } from ""express"";
import { FollowService } from ""../services/follow.service"";
import { followUserSchema } from ""../validations/follow.schema"";
import { UnauthorizedError, ValidationError } from ""@/utils/errors.utils"";

export class FollowController {
  private followService = new FollowService();

  /**
   * Follow a user
   */
  followUser = async (req: Request, res: Response, next: NextFunction) => {
    try {
      if (!req.user || !req.user.id) {
        throw new UnauthorizedError(""Not authenticated"");
      }

      const validationResult = followUserSchema.safeParse(req.body);
      if (!validationResult.success) {
        throw new ValidationError(validationResult.error.format());
      }

      const result = await this.followService.followUser(
        req.user.id,
        validationResult.data.username
      );

      res.status(200).json({
        message: ""Successfully followed user"",
        ...result,
      });
    } catch (error) {
      next(error);
    }
  };

  /**
   * Unfollow a user
   */
  unfollowUser = async (req: Request, res: Response, next: NextFunction) => {
    try {
      if (!req.user || !req.user.id) {
        throw new UnauthorizedError(""Not authenticated"");
      }

      const { username } = req.params;

      const result = await this.followService.unfollowUser(
        req.user.id,
        username
      );

      res.status(200).json({
        message: ""Successfully unfollowed user"",
        ...result,
      });
    } catch (error) {
      next(error);
    }
  };

  /**
   * Check follow status
   */
  checkFollowStatus = async (
    req: Request,
    res: Response,
    next: NextFunction
  ) => {
    try {
      if (!req.user || !req.user.id) {
        throw new UnauthorizedError(""Not authenticated"");
      }

      const { username } = req.params;

      const result = await this.followService.checkFollowStatus(
        req.user.id,
        username
      );

      res.status(200).json(result);
    } catch (error) {
      next(error);
    }
  };

  /**
   * Get followers of a user
   */
  getFollowers = async (req: Request, res: Response, next: NextFunction) => {
    try {
      const { username } = req.params;
      const page = req.query.page ? parseInt(req.query.page as string) : 1;
      const limit = req.query.limit ? parseInt(req.query.limit as string) : 20;

      const result = await this.followService.getFollowers(
        username,
        page,
        limit
      );

      res.status(200).json(result);
    } catch (error) {
      next(error);
    }
  };

  /**
   * Get users that a user is following
   */
  getFollowing = async (req: Request, res: Response, next: NextFunction) => {
    try {
      const { username } = req.params;
      const page = req.query.page ? parseInt(req.query.page as string) : 1;
      const limit = req.query.limit ? parseInt(req.query.limit as string) : 20;

      const result = await this.followService.getFollowing(
        username,
        page,
        limit
      );

      res.status(200).json(result);
    } catch (error) {
      next(error);
    }
  };

  /**
   * Get follow counts
   */
  getFollowCounts = async (req: Request, res: Response, next: NextFunction) => {
    try {
      const { username } = req.params;

      const result = await this.followService.getFollowCounts(username);

      res.status(200).json(result);
    } catch (error) {
      next(error);
    }
  };
}"

LINK NUMBER 170

File path: src/components/profile/validations/profile.schema.ts
"import { UseFormReturn } from ""react-hook-form"";
import {
  AlertDialog,
  AlertDialogContent,
  AlertDialogFooter,
  AlertDialogHeader,
  AlertDialogTitle,
} from ""../ui/alert-dialog"";
import { Button } from ""../ui/button"";
import {
  Form,
  FormControl,
  FormField,
  FormItem,
  FormLabel,
  FormMessage,
} from ""../ui/form"";
import { Input } from ""../ui/input"";
import { Separator } from ""../ui/separator"";
import { Textarea } from ""../ui/textarea"";
import { ProfileFormInput } from ""./validations/profile.schema"";

interface EditProfileForm {
  isEditProfileOpen: boolean;
  setIsEditProfileOpen: (open: boolean) => void;
  form: UseFormReturn<ProfileFormInput>;
  handleProfileUpdate: (values: ProfileFormInput) => void;
}

const EditProfileForm = ({
  isEditProfileOpen,
  setIsEditProfileOpen,
  form,
  handleProfileUpdate,
}: EditProfileForm) => {
  return (
    <AlertDialog open={isEditProfileOpen} onOpenChange={setIsEditProfileOpen}>
      <AlertDialogContent className=""max-w-md"">
        <AlertDialogHeader>
          <AlertDialogTitle>Edit Profile</AlertDialogTitle>
        </AlertDialogHeader>

        <Form {...form}>
          <form
            onSubmit={form.handleSubmit(handleProfileUpdate)}
            className=""space-y-4""
          >
            {/* <div className=""flex justify-center mb-4"">
              <div className=""relative"">
                <Avatar className=""w-24 h-24 border-2 border-background"">
                  <AvatarImage src={user.profileImageUrl} alt={user.username} />
                  <AvatarFallback className=""text-2xl bg-primary/10"">
                    {user.username.charAt(0).toUpperCase()}
                  </AvatarFallback>
                </Avatar>

                <MediaUploader type=""profile"" multiple={false} />
              </div>
            </div> */}

            <FormField
              control={form.control}
              name=""fullName""
              render={({ field }) => (
                <FormItem>
                  <FormLabel>Full Name</FormLabel>
                  <FormControl>
                    <Input placeholder=""Your full name"" {...field} />
                  </FormControl>
                  <FormMessage />
                </FormItem>
              )}
            />

            <FormField
              control={form.control}
              name=""bio""
              render={({ field }) => (
                <FormItem>
                  <FormLabel>Bio</FormLabel>
                  <FormControl>
                    <Textarea placeholder=""Tell us about yourself"" {...field} />
                  </FormControl>
                  <FormMessage />
                </FormItem>
              )}
            />

            <FormField
              control={form.control}
              name=""website""
              render={({ field }) => (
                <FormItem>
                  <FormLabel>Website</FormLabel>
                  <FormControl>
                    <Input placeholder=""https://yourwebsite.com"" {...field} />
                  </FormControl>
                  <FormMessage />
                </FormItem>
              )}
            />

            <Separator />

            <AlertDialogFooter>
              <Button
                type=""button""
                variant=""outline""
                onClick={() => setIsEditProfileOpen(false)}
              >
                Cancel
              </Button>
              <Button type=""submit"">Save Changes</Button>
            </AlertDialogFooter>
          </form>
        </Form>
      </AlertDialogContent>
    </AlertDialog>
  );
};

export default EditProfileForm;"

LINK NUMBER 171
Not enough lines

LINK NUMBER 172
Not enough lines

LINK NUMBER 173
Not enough lines

LINK NUMBER 174
Not enough lines

LINK NUMBER 175
Too many lines

LINK NUMBER 176
Not enough lines

LINK NUMBER 177

File path: src/features/employees/store/employeesSlice.ts
"import { PaginatedResponse } from ""@/api/common/commonApi.types"";
import { employeeApi } from ""@/api/employee/employeeApi"";
import {
  Employee,
  EmployeeQueryParams,
} from ""@/api/employee/employeeApi.types"";
import { addAlert, setLoading } from ""@/store/uiSlice"";
import { createAsyncThunk, createSlice } from ""@reduxjs/toolkit"";

interface EmployeeState {
  employees: PaginatedResponse<Employee> | null;
  currentEmployee: Employee | null;
  isLoading: boolean;
  error: string | null;
}

const initialState: EmployeeState = {
  employees: null,
  currentEmployee: null,
  isLoading: false,
  error: null,
};

export const getEmployees = createAsyncThunk<
  PaginatedResponse<Employee>,
  EmployeeQueryParams | undefined,
  { rejectValue: { message: string } }
>(""employees/getEmployees"", async (params, { dispatch, rejectWithValue }) => {
  try {
    dispatch(setLoading({ key: ""getEmployees"", isLoading: true }));

    return await employeeApi.getEmployees(params);
  } catch (error: any) {
    const errorMessage =
      error.response?.data?.message || ""Failed to fetch employees"";

    dispatch(addAlert({ type: ""error"", message: errorMessage }));

    return rejectWithValue({ message: errorMessage });
  } finally {
    dispatch(setLoading({ key: ""getEmployees"", isLoading: false }));
  }
});

export const getEmployeeById = createAsyncThunk<
  Employee,
  string,
  { rejectValue: { message: string } }
>(""employees/getEmployeeById"", async (id, { dispatch, rejectWithValue }) => {
  try {
    dispatch(setLoading({ key: ""getEmployeeById"", isLoading: true }));

    return employeeApi.getEmployee(id);
  } catch (error: any) {
    const errorMessage =
      error.response?.data?.message || ""Failed to fetch the employee by id"";
    dispatch(addAlert({ type: ""error"", message: errorMessage }));

    return rejectWithValue({ message: errorMessage });
  } finally {
    dispatch(setLoading({ key: ""getEmployeeById"", isLoading: false }));
  }
});

export const createEmployee = createAsyncThunk<
  Employee,
  any,
  { rejectValue: { message: string } }
>(
  ""employees/createEmployee"",
  async (employeeData, { dispatch, rejectWithValue }) => {
    try {
      dispatch(setLoading({ key: ""createEmployee"", isLoading: true }));

      const employee = await employeeApi.createEmployee(employeeData);

      dispatch(
        addAlert({ type: ""success"", message: ""Employee created successfully"" })
      );
      return employee;
    } catch (error: any) {
      const errorMessage =
        error.response?.data?.message || ""Failed to create employee"";
      dispatch(addAlert({ type: ""error"", message: errorMessage }));
      return rejectWithValue(errorMessage);
    } finally {
      dispatch(setLoading({ key: ""createEmployee"", isLoading: false }));
    }
  }
);

export const updateEmployee = createAsyncThunk<
  Employee,
  { id: string; data: any },
  { rejectValue: { message: string } }
>(
  ""employees/updateEmployee"",
  async ({ id, data }, { dispatch, rejectWithValue }) => {
    try {
      dispatch(setLoading({ key: ""updateEmployee"", isLoading: true }));

      const employee = await employeeApi.updateEmployee(id, data);

      dispatch(
        addAlert({ type: ""success"", message: ""Employee updated successfully"" })
      );

      return employee;
    } catch (error: any) {
      const errorMessage =
        error.response?.data?.message || ""Failed to update employee"";
      dispatch(addAlert({ type: ""error"", message: errorMessage }));

      return rejectWithValue({ message: errorMessage });
    } finally {
      dispatch(setLoading({ key: ""updateEmployee"", isLoading: false }));
    }
  }
);

export const deleteEmployee = createAsyncThunk<
  string,
  string,
  { rejectValue: { message: string } }
>(""employees/deleteEmployee"", async (id, { dispatch, rejectWithValue }) => {
  try {
    dispatch(setLoading({ key: ""deleteEmployee"", isLoading: true }));

    await employeeApi.deleteEmployee(id);

    dispatch(
      addAlert({ type: ""success"", message: ""Employee deleted successfully"" })
    );

    return id;
  } catch (error: any) {
    const errorMessage =
      error.response?.data?.message || ""Failed to delete employee"";
    dispatch(addAlert({ type: ""error"", message: errorMessage }));
    return rejectWithValue({ message: errorMessage });
  } finally {
    dispatch(setLoading({ key: ""deleteEmployee"", isLoading: false }));
  }
});

const employeesSlice = createSlice({
  name: ""employees"",
  initialState,
  reducers: {
    clearEmployeeError: (state) => {
      state.error = null;
    },
    clearCurrentEmployee: (state) => {
      state.currentEmployee = null;
    },
  },
  extraReducers: (builder) => {
    builder
      // Get Employees
      .addCase(getEmployees.pending, (state) => {
        state.isLoading = true;
        state.error = null;
      })
      .addCase(getEmployees.fulfilled, (state, action) => {
        state.isLoading = false;
        state.employees = action.payload;
        state.error = null;
      })
      .addCase(getEmployees.rejected, (state, action) => {
        state.isLoading = false;
        state.error = action.payload?.message || ""Failed to fetch employees"";
      })
      // Get Employee by Id
      .addCase(getEmployeeById.pending, (state) => {
        state.isLoading = true;
        state.error = null;
      })
      .addCase(getEmployeeById.fulfilled, (state, action) => {
        state.isLoading = false;
        state.currentEmployee = action.payload;
        state.error = null;
      })
      .addCase(getEmployeeById.rejected, (state, action) => {
        state.isLoading = false;
        state.error = action.payload?.message || ""Failed to fetch employee"";
      })

      // createEmployee
      .addCase(createEmployee.fulfilled, (state) => {
        state.isLoading = false;
      })

      // updateEmployee
      .addCase(updateEmployee.fulfilled, (state, action) => {
        state.isLoading = false;
        state.currentEmployee = action.payload;
      })

      // deleteEmployee
      .addCase(deleteEmployee.fulfilled, (state, action) => {
        state.isLoading = false;
        if (state.employees && state.employees.data) {
          state.employees.data = state.employees.data.filter(
            (employee) => employee.id !== action.payload
          );
        }
      });
  },
});

export const { clearEmployeeError, clearCurrentEmployee } =
  employeesSlice.actions;
export default employeesSlice.reducer;"

LINK NUMBER 178
Not enough lines

LINK NUMBER 179
Not enough lines

LINK NUMBER 180
Not enough lines

LINK NUMBER 181

File path: src/features/follow/validations/follow.schema.ts
"import { AppDataSource } from ""@/config/database"";
import { Follow } from ""@/entities/Follow.entity"";
import { User } from ""@/entities/User.entity"";
import { ForbiddenError, NotFoundError } from ""@/utils/errors.utils"";

export class FollowService {
  private followRepository = AppDataSource.getRepository(Follow);
  private userRepository = AppDataSource.getRepository(User);

  /**
   * Follow user
   */
  async followUser(followerUserId: string, usernameToFollow: string) {
    // Find follower
    const follower = await this.userRepository.findOneBy({
      id: followerUserId,
    });
    if (!follower) {
      throw new NotFoundError(""User not found"");
    }

    // Find user to follow
    const userToFollow = await this.userRepository.findOneBy({
      username: usernameToFollow,
    });
    if (!userToFollow) {
      throw new NotFoundError(""User to follow not found"");
    }

    // Cannot follow yourself
    if (follower.id === userToFollow.id) {
      throw new ForbiddenError(""You cannot follow yourself"");
    }

    // Check if already following
    const existingFollow = await this.followRepository.findOne({
      where: {
        follower: { id: followerUserId },
        following: { id: userToFollow.id },
      },
    });

    if (existingFollow) {
      throw new ForbiddenError(""You are already following this user"");
    }

    // Create follow relation
    const follow = this.followRepository.create({
      follower,
      following: userToFollow,
    });

    await this.followRepository.save(follow);

    return { success: true };
  }

  /**
   * Unfollow a user
   */
  async unfollowUser(followerUserId: string, usernameToUnfollow: string) {
    // Find user to unfollow
    const userToUnfollow = await this.userRepository.findOneBy({
      username: usernameToUnfollow,
    });
    if (!userToUnfollow) {
      throw new NotFoundError(""User to unfollow not found"");
    }

    // Check if following
    const follow = await this.followRepository.findOne({
      where: {
        follower: { id: followerUserId },
        following: { id: userToUnfollow.id },
      },
    });

    if (!follow) {
      throw new ForbiddenError(""You are not following this user"");
    }

    // Remove follow relationship
    await this.followRepository.remove(follow);

    return { success: true };
  }

  /**
   * Check if one user follows another
   */
  async checkFollowStatus(followerUserId: string, usernameToCheck: string) {
    const userToCheck = await this.userRepository.findOneBy({
      username: usernameToCheck,
    });
    if (!userToCheck) {
      throw new NotFoundError(""User not found"");
    }

    const follow = await this.followRepository.findOne({
      where: {
        follower: { id: followerUserId },
        following: { id: userToCheck.id },
      },
    });

    return { following: !!follow };
  }

  /**
   * Get followers of a user
   */
  async getFollowers(username: string, page = 1, limit = 20) {
    const user = await this.userRepository.findOneBy({ username });
    if (!user) {
      throw new NotFoundError(""User not found"");
    }

    const skip = (page - 1) * limit;

    const [follows, total] = await this.followRepository
      .createQueryBuilder(""follow"")
      .leftJoinAndSelect(""follow.follower"", ""follower"")
      .where(""follow.following.id = :userId"", { userId: user.id })
      .skip(skip)
      .take(limit)
      .orderBy(""follow.createdAt"", ""DESC"")
      .getManyAndCount();

    // Extract follower users only
    const followers = follows.map((follow) => follow.follower);

    return {
      followers,
      meta: {
        total,
        page,
        limit,
        totalPages: Math.ceil(total / limit),
      },
    };
  }

  /**
   * Get users that a user is following
   */
  async getFollowing(username: string, page = 1, limit = 20) {
    const user = await this.userRepository.findOneBy({ username });
    if (!user) {
      throw new NotFoundError(""User not found"");
    }

    const skip = (page - 1) * limit;

    const [follows, total] = await this.followRepository
      .createQueryBuilder(""follow"")
      .leftJoinAndSelect(""follow.following"", ""following"")
      .where(""follow.follower.id = :userId"", { userId: user.id })
      .skip(skip)
      .take(limit)
      .orderBy(""follow.createdAt"", ""DESC"")
      .getManyAndCount();

    // Extract following users only
    const following = follows.map((follow) => follow.following);

    return {
      following,
      meta: {
        total,
        page,
        limit,
        totalPages: Math.ceil(total / limit),
      },
    };
  }

  /**
   * Get follower/following counts
   */
  async getFollowCounts(username: string) {
    const user = await this.userRepository.findOneBy({ username });
    if (!user) {
      throw new NotFoundError(""User not found"");
    }

    const followersCount = await this.followRepository.count({
      where: { following: { id: user.id } },
    });

    const followingCount = await this.followRepository.count({
      where: { follower: { id: user.id } },
    });

    return {
      username,
      followersCount,
      followingCount,
    };
  }
}"

LINK NUMBER 182